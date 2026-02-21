from collections import defaultdict
from datetime import date, datetime
from decimal import Decimal
import json
from django.contrib import messages
from django.db.models import Sum, F, DecimalField, Q
from django.shortcuts import render, get_object_or_404, redirect
from django.utils import timezone
from urllib.parse import quote_plus
from django.utils.dateparse import parse_date
from django.forms import modelformset_factory
from django.db.models.functions import TruncMonth
from autocarros.decorators import acesso_restrito
from .models import Autocarro, CobradorViagem, Comprovativo, ComprovativoRelatorio, Deposito, Despesa2, DespesaCombustivel, DespesaFixa, Manutencao, RegistoDiario, Despesa, RegistroKM, RegistroKMItem, RelatorioSector, Sector, Motorista, SubCategoriaDespesa
from .forms import DespesaCombustivelForm, DespesaFixaForm, DespesaForm2, EstadoAutocarroForm, AutocarroForm, DespesaForm, ComprovativoFormSet, ManutencaoForm, MultiFileForm,RegistoDiarioFormSet, RelatorioSectorForm, SectorForm, SectorGestorForm, SelecionarSectorCombustivelForm, RegistoDiarioForm, SubCategoriaDespesaForm
from autocarros import models
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.views import View
from .forms import CustomUserCreationForm, CustomAuthenticationForm, UserUpdateForm
from .models import CustomUser
from django.utils.text import slugify
from django.contrib.auth.models import User
import json
from django.views.decorators.http import require_POST




# === Decorator para só admins poderem associar gestores === #
@login_required
@acesso_restrito(['admin'])
def admin_required(user):
    return user.is_authenticated and user.is_admin()


@login_required
@acesso_restrito(['admin'])
def associar_gestor(request, sector_id):
    sector = get_object_or_404(Sector, id=sector_id)

    if request.method == "POST":
        form = SectorGestorForm(request.POST, instance=sector)
        if form.is_valid():
            form.save()
            messages.success(request, f'Gestor associado ao setor "{sector.nome}" com sucesso!')
            return redirect('lista_sectores')
    else:
        form = SectorGestorForm(instance=sector)

    return render(request, "autocarros/associar_gestor.html", {
        "form": form,
        "sector": sector,
    })


@login_required
@acesso_restrito(['admin'])
def admin_required(view_func):
    decorated_view = user_passes_test(
        lambda user: user.is_authenticated and user.is_admin(),
        login_url='acesso_negado'
    )(view_func)
    return decorated_view


@login_required
@acesso_restrito(['admin'])
def gestor_required(view_func):
    decorated_view = user_passes_test(
        lambda user: user.is_authenticated and user.is_gestor(),
        login_url='acesso_negado'
    )(view_func)
    return decorated_view


@login_required
def can_edit_required(view_func):
    decorated_view = user_passes_test(
        lambda user: user.is_authenticated and user.can_edit(),
        login_url='acesso_negado'
    )(view_func)
    return decorated_view

# === Login View === #
class LoginView(View):
    def get(self, request):
        return render(request, 'auth/login.html')

    def post(self, request):
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            messages.success(request, f"Bem-vindo(a), {user.username}!")
            return redirect('dashboard')
        else:
            messages.error(request, "Usuário ou senha incorretos.")
            return render(request, 'auth/login.html')


@login_required
@acesso_restrito(['admin'])
def register_user(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')
        telefone = request.POST.get('telefone')
        password = request.POST.get('password')
        confirm = request.POST.get('confirm')
        nivel = request.POST.get('nivel_acesso')

        if password != confirm:
            messages.error(request, "As senhas não coincidem.")
            return redirect('register')

        if CustomUser.objects.filter(username=username).exists():
            messages.error(request, "Nome de usuário já existe.")
            return redirect('register')

        user = CustomUser.objects.create_user(
            username=username,
            email=email,
            password=password,
            telefone=telefone,
            nivel_acesso=nivel
        )

        messages.success(request, f"Usuário '{user.username}' criado com sucesso!")
        return redirect('dashboard')  # ou a tua página principal

    niveis = CustomUser.NIVEL_ACESSO_CHOICES
    return render(request, 'auth/register.html', {'niveis': niveis})


@login_required
def logout_view(request):
    logout(request)
    return redirect('login')


# === Admin Dashboard View === #
@login_required
@acesso_restrito(['admin'])
def admin_dashboard(request):
    usuarios = CustomUser.objects.all()
    context = {
        'usuarios': usuarios,
        'total_usuarios': usuarios.count(),
        'usuarios_ativos': usuarios.filter(ativo=True).count(),
    }
    return render(request, 'admin_dashboard.html', context)


#=== Perfil do Usuário View === #
@login_required
def perfil(request):
    return render(request, 'autocarros/perfil.html', {'user': request.user})


@login_required
@acesso_restrito(['admin'])
def gerir_usuarios(request):
    usuarios = CustomUser.objects.all().order_by('-date_joined')
    return render(request, 'gerir_usuarios.html', {'usuarios': usuarios})


@login_required
@acesso_restrito(['admin'])
def editar_usuario(request, user_id):
    usuario = get_object_or_404(CustomUser, id=user_id)
    if request.method == 'POST':
        form = UserUpdateForm(request.POST, instance=usuario)
        if form.is_valid():
            form.save()
            messages.success(request, f'Usuário {usuario.username} atualizado com sucesso!')
            return redirect('gerir_usuarios')
    else:
        form = UserUpdateForm(instance=usuario)
    
    return render(request, 'editar_usuario.html', {'form': form, 'usuario': usuario})


# === Acesso Negado View === #
def acesso_negado(request):
    return render(request, 'acesso_negado.html', status=403)


# === Verificar Integridade dos Dados === #
@login_required
@acesso_restrito(['admin'])
def verificar_integridade(request):
    """View para verificar e corrigir problemas de integridade"""
    if not request.user.is_superuser:
        messages.error(request, "❌ Apenas administradores podem acessar esta função.")
        return redirect('listar_registros')
    
    problemas = []
    
    # 🔹 VERIFICAR REGISTROS SEM RELATÓRIO
    registros_sem_relatorio = RegistoDiario.objects.filter(relatorio__isnull=True)
    if registros_sem_relatorio.exists():
        problemas.append(f"❌ {registros_sem_relatorio.count()} registros sem relatório associado")
    
    # 🔹 VERIFICAR RELATÓRIOS SEM REGISTROS
    relatorios_sem_registros = RelatorioSector.objects.filter(registos__isnull=True)
    if relatorios_sem_registros.exists():
        problemas.append(f"❌ {relatorios_sem_registros.count()} relatórios sem registros")
    
    # 🔹 VERIFICAR DUPLICAÇÕES
    from django.db.models import Count
    duplicatas = RegistoDiario.objects.values('relatorio', 'autocarro').annotate(
        count=Count('id')
    ).filter(count__gt=1)
    
    if duplicatas.exists():
        problemas.append(f"❌ {duplicatas.count()} duplicatas encontradas")
    
    if request.method == "POST" and "corrigir" in request.POST:
        # 🔹 CORRIGIR REGISTROS SEM RELATÓRIO
        for registro in registros_sem_relatorio:
            relatorio_compativel = RelatorioSector.objects.filter(
                sector=registro.autocarro.sector,
                data=registro.data
            ).first()
            
            if relatorio_compativel:
                registro.relatorio = relatorio_compativel
                registro.save()
        
        # 🔹 CRIAR REGISTROS PARA RELATÓRIOS VAZIOS
        for relatorio in relatorios_sem_registros:
            autocarros = Autocarro.objects.filter(sector=relatorio.sector)
            for autocarro in autocarros:
                RegistoDiario.objects.get_or_create(
                    relatorio=relatorio,
                    autocarro=autocarro,
                    defaults={'data': relatorio.data}
                )
        
        messages.success(request, "✅ Problemas de integridade corrigidos!")
        return redirect('verificar_integridade')
    
    context = {
        'problemas': problemas,
        'total_problemas': len(problemas),
    }
    return render(request, "autocarros/verificar_integridade.html", context)


# === Sector Management Views === #
@login_required
def layout_base(request):
    sectores = Sector.objects.all()
    return render(request, "base.html", {"sectores": sectores})


# === Lista de Setores === #
@login_required
@acesso_restrito(['admin'])
def lista_sectores(request):
    sectores = Sector.objects.all().order_by("nome")
    return render(request, "autocarros/lista_sectores.html", {"sectores": sectores})


@login_required
@acesso_restrito(['admin'])
def adicionar_sector(request):
    if request.method == "POST":
        form = SectorForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Setor adicionado com sucesso!")
            return redirect("lista_sectores")
        else:
            messages.error(request, "❌ Erro ao adicionar setor. Verifique os dados.")
    else:
        form = SectorForm()
    return render(request, "autocarros/adicionar_sector.html", {"form": form})


@login_required
@acesso_restrito(['admin'])
def editar_sector(request, pk):
    sector = get_object_or_404(Sector, pk=pk)
    if request.method == "POST":
        form = SectorForm(request.POST, instance=sector)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Setor atualizado com sucesso!")
            return redirect("lista_sectores")
        else:
            messages.error(request, "❌ Erro ao atualizar setor. Verifique os dados.")
    else:
        form = SectorForm(instance=sector)
    return render(request, "autocarros/adicionar_sector.html", {"form": form, "editar": True})


@login_required
@acesso_restrito(['admin'])
def apagar_sector(request, pk):
    sector = get_object_or_404(Sector, pk=pk)
    if request.method == "POST":
        try:
            sector.delete()
            messages.success(request, "✅ Setor apagado com sucesso!")
            return redirect("lista_sectores")
        except Exception as e:
            messages.error(request, f"❌ Erro ao apagar setor: {str(e)}")
    return render(request, "autocarros/confirmar_apagar_sector.html", {"sector": sector})


# === Dashboard View === #
@login_required
@acesso_restrito(['admin'])
def dashboard(request):
    hoje = timezone.now().date()

    # 🔹 Capturar "YYYY-MM" vindo do input type="month"
    mes_param = request.GET.get("mes", hoje.strftime("%Y-%m"))
    try:
        ano, mes = map(int, mes_param.split("-"))
    except ValueError:
        ano, mes = hoje.year, hoje.month

    # 🔹 anos disponíveis
    anos_disponiveis = [
        int(d.year) for d in RegistoDiario.objects.dates("data", "year", order="DESC")
    ]
    if hoje.year not in anos_disponiveis:
        anos_disponiveis.insert(0, hoje.year)

    # 🔹 registos filtrados
    registos = RegistoDiario.objects.filter(
        data__year=ano,
        data__month=mes
    ).select_related("autocarro")

    # 🔹 Combustível por autocarro/data
    combustivel_map_dashboard = {}
    if registos.exists():
        autocarro_ids = set(registos.values_list('autocarro_id', flat=True))
        datas = set(registos.values_list('data', flat=True))

        combustiveis_dash = DespesaCombustivel.objects.filter(
            autocarro_id__in=autocarro_ids,
            data__in=datas
        )

        agg_dash = defaultdict(lambda: {
            'total_valor': Decimal('0'),
            'total_valor_litros': Decimal('0'),
            'total_sobragem': Decimal('0'),
            'total_lavagem': Decimal('0')
        })

        for c in combustiveis_dash:
            key = f"{c.autocarro_id}_{c.data.isoformat()}"
            agg_dash[key]['total_valor'] += c.valor or Decimal('0')
            agg_dash[key]['total_valor_litros'] += c.valor_litros or Decimal('0')
            agg_dash[key]['total_sobragem'] += c.sobragem_filtros or Decimal('0')
            agg_dash[key]['total_lavagem'] += c.lavagem or Decimal('0')

        combustivel_map_dashboard.update(agg_dash)

    # 🔹 Totais gerais
    total_entradas = registos.aggregate(
        total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
    )["total"] or Decimal("0")

    total_saidas_registos = registos.aggregate(
        total=Sum(
            F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"),
            output_field=DecimalField()
        )
    )["total"] or Decimal("0")

    total_saidas_despesas = Despesa.objects.filter(
        data__year=ano,
        data__month=mes
    ).aggregate(
        total=Sum("valor", output_field=DecimalField())
    )["total"] or Decimal("0")

    # 🔹 Despesa geral dos setores (NOVO)
    total_despesa_geral = RelatorioSector.objects.filter(
        data__year=ano,
        data__month=mes
    ).aggregate(
        total=Sum('despesa_geral', output_field=DecimalField())
    )['total'] or Decimal('0')

    total_alimentacao_estaleiro = RelatorioSector.objects.filter(
        data__year=ano,
        data__month=mes
    ).aggregate(
        total=Sum('alimentacao_estaleiro', output_field=DecimalField())
    )['total'] or Decimal('0')

    qs_fixas = DespesaFixa.objects.filter(ativo=True)

    # Mensais: contam a partir do mês de início até o presente mês
    mensais_qs = qs_fixas.filter(
        periodicidade__iexact='mensal',
        data_inicio__lte=date(ano, mes, 1)
    )

    # Anuais: contam apenas se data_inicio for no mesmo mês/ano
    anuais_qs = qs_fixas.filter(
        periodicidade__iexact='anual',
        data_inicio__year=ano,
        data_inicio__month=mes
    )

    # Únicas: contam apenas no mês/ano específico
    unicas_qs = qs_fixas.filter(
        periodicidade__iexact='único',
        data_inicio__year=ano,
        data_inicio__month=mes
    )

    total_despesas_fixas = (
        mensais_qs.aggregate(total=Sum('valor', output_field=DecimalField()))['total'] or Decimal('0')
    ) + (
        anuais_qs.aggregate(total=Sum('valor', output_field=DecimalField()))['total'] or Decimal('0')
    ) + (
        unicas_qs.aggregate(total=Sum('valor', output_field=DecimalField()))['total'] or Decimal('0')
    )

    total_combustivel = DespesaCombustivel.objects.filter(
        data__year=ano,
        data__month=mes
    ).aggregate(
        total_valor=Sum('valor', output_field=DecimalField()),
        total_litros=Sum('valor_litros', output_field=DecimalField()),
        total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
        total_lavagem=Sum('lavagem', output_field=DecimalField()),
    )

    total_combustivel_valor = total_combustivel.get('total_valor') or Decimal('0')
    total_combustivel_sobragem = total_combustivel.get('total_sobragem') or Decimal('0')
    total_combustivel_lavagem = total_combustivel.get('total_lavagem') or Decimal('0')

    # 🔹 Total Despesa2 (variáveis do mês)
    total_despesa2 = Despesa2.objects.filter(
        data__year=ano,
        data__month=mes
    ).aggregate(
        total=Sum("valor", output_field=DecimalField())
    )["total"] or Decimal("0")

    # total de saídas inclui registos, despesas (Despesas) e combustíveis + sobragem/lavagem + despesas fixas
    total_saidas = (
            total_saidas_registos
            + total_combustivel_valor
            + total_combustivel_sobragem
            + total_combustivel_lavagem
            + total_despesa_geral
            + total_alimentacao_estaleiro
    )
    total_resto = total_entradas - total_saidas
    total_despesa2_1 = total_despesa2 + total_despesas_fixas + total_saidas_despesas
    total_lucro = total_resto - total_despesa2_1

    # 🔹 Estatísticas por autocarro
    autocarros_stats = []
    for autocarro in Autocarro.objects.all():
        registos_auto = registos.filter(autocarro=autocarro)
        stats = {
            "autocarro": autocarro,
            "total_km": registos_auto.aggregate(Sum("km_percorridos"))["km_percorridos__sum"] or 0,
            "total_entradas": registos_auto.aggregate(
                total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
            )["total"] or Decimal('0'),
            "total_saidas": registos_auto.aggregate(
                total=Sum(F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"), output_field=DecimalField())
            )["total"] or Decimal('0'),
            "total_passageiros": registos_auto.aggregate(Sum("numero_passageiros"))["numero_passageiros__sum"] or 0,
            "total_viagens": registos_auto.aggregate(Sum("numero_viagens"))["numero_viagens__sum"] or 0,
        }

        comb_auto = DespesaCombustivel.objects.filter(
            autocarro=autocarro,
            data__year=ano,
            data__month=mes
        ).aggregate(
            total_valor=Sum('valor', output_field=DecimalField()),
            total_litros=Sum('valor_litros', output_field=DecimalField()),
            total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
            total_lavagem=Sum('lavagem', output_field=DecimalField()),
        )

        comb_val = comb_auto.get('total_valor') or Decimal('0')
        stats['total_combustivel'] = comb_val

        # em vez de 'litros' o ficheiro pede 'alimentacao + outros' por autocarro
        alim_outros_auto = registos_auto.aggregate(
            total=Sum(F("alimentacao") + F("outros"), output_field=DecimalField())
        )["total"] or Decimal('0')

        stats['total_alim_outros'] = alim_outros_auto

        stats['total_combustivel_litros'] = comb_auto.get('total_litros') or Decimal('0')
        stats['total_combustivel_sobragem'] = comb_auto.get('total_sobragem') or Decimal('0')

        stats['total_combustivel_lavagem'] = comb_auto.get('total_lavagem') or Decimal('0')

        comb_sobr = stats['total_combustivel_sobragem']

        comb_lav = stats['total_combustivel_lavagem']

        # incluir combustível e respetivas taxas nas saídas por autocarro
        stats['total_saidas'] += stats['total_combustivel'] + comb_sobr + comb_lav
        # OBS: 'total_alim_outros' já faz parte de 'total_saidas' (porque veio de registos_auto agregados),
        # mas mantemos o campo separado para exibição no lugar de "litros".
        stats["resto"] = stats["total_entradas"] - stats["total_saidas"]

        autocarros_stats.append(stats)

    # 🔹 Calcular o maior saldo (mais lucrativo)
    max_saldo = max((a["resto"] for a in autocarros_stats), default=Decimal('0'))

    # 🔹 Registos recentes
    registos_recentes_qs = registos.order_by("-data")[:10]
    registos_recentes = []
    for reg in registos_recentes_qs:
        key = f"{reg.autocarro_id}_{reg.data.isoformat()}"
        comb = combustivel_map_dashboard.get(key, {})
        reg.combustivel_total = comb.get('total_valor', Decimal('0'))
        # em vez de litros, mostramos alimentacao + outros do próprio registo
        reg.alim_outros = (getattr(reg, 'alimentacao', Decimal('0')) or Decimal('0')) + (getattr(reg, 'outros', Decimal('0')) or Decimal('0'))
        reg.combustivel_valor_litros = comb.get('total_valor_litros', Decimal('0'))
        reg.combustivel_sobragem = comb.get('total_sobragem', Decimal('0'))
        reg.combustivel_lavagem = comb.get('total_lavagem', Decimal('0'))
        reg.saidas_total_incl_combustivel = (
            reg.saidas_total() + reg.combustivel_total + reg.combustivel_sobragem + reg.combustivel_lavagem
        )
        reg.saldo_liquido_incl_combustivel = reg.entradas_total() - reg.saidas_total_incl_combustivel
        registos_recentes.append(reg)


    context = {
        "ano": ano,
        "mes": f"{ano}-{mes:02d}",
        "anos_disponiveis": anos_disponiveis,
        "total_entradas": total_entradas,
        "total_saidas": total_saidas,
        "total_saidas_registos": total_saidas_registos,
        "total_despesa_geral": total_despesa_geral,
        "total_saidas_despesas": total_saidas_despesas,
        "total_despesa2_1": total_despesa2_1,
        "total_despesas_fixas": total_despesas_fixas,
        "total_resto": total_resto,
        "total_lucro": total_lucro,
        "autocarros_stats": autocarros_stats,
        "registos_recentes": registos_recentes,
        "max_saldo": max_saldo,
    }
    return render(request, "autocarros/dashboard.html", context)


#================================== Arquivo World ========================================== #
from django.http import HttpResponse, HttpResponseBadRequest, JsonResponse
from django.utils import timezone
from django.db.models import Sum, F, DecimalField
from django.contrib.humanize.templatetags.humanize import intcomma
from decimal import Decimal
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
from .models import RegistoDiario, Despesa, DespesaCombustivel, Autocarro
from .decorators import acesso_restrito
from django.contrib.auth.decorators import login_required

# === Exportar Relatório do Dashboard === #
@login_required
@acesso_restrito(['admin'])
def exportar_relatorio_dashboard(request):
    hoje = timezone.now().date()
    mes_param = request.GET.get("mes", hoje.strftime("%Y-%m"))

    try:
        ano, mes = map(int, mes_param.split("-"))
    except ValueError:
        ano, mes = hoje.year, hoje.month

    registos = RegistoDiario.objects.filter(data__year=ano, data__month=mes)

    total_entradas = registos.aggregate(
        total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
    )["total"] or Decimal("0")

    total_saidas_registos = registos.aggregate(
        total=Sum(F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"), output_field=DecimalField())
    )["total"] or Decimal("0")

    total_saidas_despesas = Despesa.objects.filter(
        data__year=ano, data__month=mes
    ).aggregate(total=Sum("valor", output_field=DecimalField()))["total"] or Decimal("0")

    # 🔹 Corrigido: campo 'valor_litros' em vez de 'litros'
    total_combustivel = DespesaCombustivel.objects.filter(
        data__year=ano, data__month=mes
    ).aggregate(
        total_valor=Sum('valor', output_field=DecimalField()),
        total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
        total_lavagem=Sum('lavagem', output_field=DecimalField()),
        total_litros=Sum('valor_litros', output_field=DecimalField())
    )

    total_combustivel_valor = total_combustivel.get('total_valor') or Decimal('0')
    total_combustivel_sobragem = total_combustivel.get('total_sobragem') or Decimal('0')
    total_combustivel_lavagem = total_combustivel.get('total_lavagem') or Decimal('0')

    total_saidas = (
        total_saidas_registos
        + total_saidas_despesas
        + total_combustivel_valor
        + total_combustivel_sobragem
        + total_combustivel_lavagem
    )
    total_resto = total_entradas - total_saidas

    # Estatísticas por autocarro
    autocarros_stats = []
    for autocarro in Autocarro.objects.all():
        registos_auto = registos.filter(autocarro=autocarro)
        if not registos_auto.exists():
            continue

        total_km = registos_auto.aggregate(Sum("km_percorridos"))["km_percorridos__sum"] or 0
        total_entr_auto = registos_auto.aggregate(
            total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
        )["total"] or Decimal("0")
        total_saida_auto = registos_auto.aggregate(
            total=Sum(F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"), output_field=DecimalField())
        )["total"] or Decimal("0")
        total_passageiros = registos_auto.aggregate(Sum("numero_passageiros"))["numero_passageiros__sum"] or 0
        total_viagens = registos_auto.aggregate(Sum("numero_viagens"))["numero_viagens__sum"] or 0

        # 🔹 Corrigido aqui também
        comb_auto = DespesaCombustivel.objects.filter(
            autocarro=autocarro, data__year=ano, data__month=mes
        ).aggregate(
            total_valor=Sum('valor', output_field=DecimalField()),
            total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
            total_lavagem=Sum('lavagem', output_field=DecimalField()),
            total_litros=Sum('valor_litros', output_field=DecimalField())
        )

        comb_val = comb_auto.get('total_valor') or Decimal('0')
        comb_sobr = comb_auto.get('total_sobragem') or Decimal('0')
        comb_lav = comb_auto.get('total_lavagem') or Decimal('0')
        comb_litros = comb_auto.get('total_litros') or Decimal('0')

        stats = {
            "autocarro": autocarro,
            "total_km": total_km,
            "total_entradas": total_entr_auto,
            "total_saidas": total_saida_auto,
            "total_passageiros": total_passageiros,
            "total_viagens": total_viagens,
            "total_combustivel": comb_val,
            "total_combustivel_litros": comb_litros,
            "total_combustivel_sobragem": comb_sobr,
            "total_combustivel_lavagem": comb_lav,
        }

        # incluir combustível nas saídas por autocarro
        stats['total_saidas'] = stats['total_saidas'] + stats.get('total_combustivel', Decimal('0')) + comb_sobr + comb_lav
        stats["resto"] = stats["total_entradas"] - stats['total_saidas']
        autocarros_stats.append(stats)

    # Criar documento Word
    doc = Document()

    # Configurar página para paisagem
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = section.right_margin = Inches(0.5)
    section.top_margin = section.bottom_margin = Inches(0.5)

    # Cabeçalho profissional
    header_table = doc.add_table(rows=1, cols=3)
    header_table.autofit = True

    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_run = left_para.add_run("🚌")
    left_run.font.size = Pt(28)
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    center_cell = header_table.cell(0, 1)
    center_para = center_cell.paragraphs[0]
    center_run = center_para.add_run(f"RELATÓRIO MENSAL - {mes_param}")
    center_run.font.size = Pt(18)
    center_run.font.bold = True
    center_run.font.color.rgb = RGBColor(13, 27, 42)
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    right_cell = header_table.cell(0, 2)
    right_para = right_cell.paragraphs[0]
    right_run = right_para.add_run(f"{datetime.now().strftime('%d/%m/%Y')}")
    right_run.font.size = Pt(10)
    right_run.font.color.rgb = RGBColor(100, 100, 100)
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph().add_run().add_break()

    # --- RESUMO GERAL ---
    h = doc.add_heading("RESUMO GERAL DO MÊS", level=2)
    try:
        h.runs[0].font.color.rgb = RGBColor(27, 42, 73)
    except Exception:
        pass

    tabela_resumo = doc.add_table(rows=2, cols=4)
    tabela_resumo.style = "Table Grid"

    # ajustar para mostrar centimos (2 casas decimais) usando babel
    from babel.numbers import format_currency

    def fmt_money_simple(x):
        try:
            val = Decimal(x or 0).quantize(Decimal('0.01'))
            return format_currency(float(val), "AOA", locale="pt_PT").replace("AOA", "Kz")
        except Exception:
            try:
                return f"{Decimal(x or 0):.2f} Kz"
            except Exception:
                return f"{x} Kz"

    entradas_val = fmt_money_simple(total_entradas)
    despesas_val = fmt_money_simple(total_saidas)
    resto_val = fmt_money_simple(total_resto)
    eficiencia_val = f"{(float(total_resto) / float(total_entradas) * 100) if total_entradas > 0 else 0:.1f}%"

    # label alterada para "SALDO"
    cards_data = [
        ("ENTRADAS TOTAIS", entradas_val, "1B4F72"),
        ("DESPESAS TOTAIS", despesas_val, "C0392B"),
        ("SALDO", resto_val, "27AE60"),
        ("EFICIÊNCIA", eficiencia_val, "8E44AD")
    ]

    for i, (titulo, valor, cor) in enumerate(cards_data):
        cell = tabela_resumo.cell(0, i)
        cell.text = titulo
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        try:
            cell._element.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:fill="{cor}"/>')
            )
        except Exception:
            pass

        cell_valor = tabela_resumo.cell(1, i)
        cell_valor.text = valor
        cell_valor.paragraphs[0].runs[0].font.bold = True
        cell_valor.paragraphs[0].runs[0].font.size = Pt(12)
        cell_valor.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().add_run().add_break()

    # --- DESPESAS ESPECÍFICAS ---
    doc.add_heading("DESPESAS OPERACIONAIS", level=2)

    despesas_data = [
        ("Combustível", total_combustivel_valor),
        ("Sopragem de Filtros", total_combustivel_sobragem),
        ("Lavagem", total_combustivel_lavagem),
        ("Outras Despesas", total_saidas_despesas + total_saidas_registos),
    ]

    # +1 linha para o cabeçalho
    tabela_despesas = doc.add_table(rows=len(despesas_data) + 1, cols=3)
    tabela_despesas.style = "Table Grid"

    # Cabeçalho
    cabecalho_despesas = ["CATEGORIA", "VALOR", "% DO TOTAL"]
    for i, titulo in enumerate(cabecalho_despesas):
        cell = tabela_despesas.cell(0, i)
        cell.text = titulo
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        try:
            cell._element.get_or_add_tcPr().append(
                parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
            )
        except Exception:
            pass

    # Função para formatar com separador de milhar e vírgula nos centavos
    def fmt_money(valor):
        try:
            valor = Decimal(valor or 0).quantize(Decimal('0.01'))
            return format_currency(float(valor), "AOA", locale="pt_PT").replace("AOA", "Kz")
        except Exception:
            try:
                return f"{Decimal(valor or 0):.2f} Kz"
            except Exception:
                return f"{valor} Kz"

    # Preenchimento das linhas da tabela
    for i, (categoria, valor) in enumerate(despesas_data, start=1):
        cell_cat = tabela_despesas.cell(i, 0)
        cell_cat.text = categoria

        cell_val = tabela_despesas.cell(i, 1)
        cell_val.text = fmt_money(valor)
        cell_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        try:
            pct = (float(valor) / float(total_saidas) * 100) if total_saidas > 0 else 0
        except Exception:
            pct = 0
        cell_pct = tabela_despesas.cell(i, 2)
        cell_pct.text = f"{pct:.1f}%"
        cell_pct.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Zebra shading (linhas alternadas)
        if i % 2 == 0:
            for c in (cell_cat, cell_val, cell_pct):
                try:
                    c._element.get_or_add_tcPr().append(
                        parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9F9"/>')
                    )
                except Exception:
                    pass

    doc.add_paragraph().add_run().add_break()

    # --- DETALHE POR AUTOCARRO ---
    doc.add_heading("DETALHE POR AUTOCARRO", level=2)

    if autocarros_stats:
        cols = ["AUTOCARRO", "KM", "ENTRADAS", "SAÍDAS", "COMBUSTÍVEL", "LITROS", "RESTO"]
        tabela_autos = doc.add_table(rows=len(autocarros_stats) + 1, cols=len(cols))
        tabela_autos.style = "Table Grid"

        # Cabeçalho
        for j, titulo in enumerate(cols):
            cell = tabela_autos.cell(0, j)
            cell.text = titulo
            cell.paragraphs[0].runs[0].font.bold = True
            try:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
                )
            except Exception:
                pass

        # Linhas
        for i, s in enumerate(autocarros_stats, start=1):
            tabela_autos.cell(i, 0).text = str(s["autocarro"].numero)
            tabela_autos.cell(i, 1).text = str(int(s.get("total_km", 0) or 0))
            tabela_autos.cell(i, 2).text = fmt_money(s.get("total_entradas", Decimal('0')))
            tabela_autos.cell(i, 3).text = fmt_money(s.get("total_saidas", Decimal('0')))
            tabela_autos.cell(i, 4).text = fmt_money(s.get("total_combustivel", Decimal('0')))
            # litros podem ser None
            litros = s.get("total_combustivel_litros", Decimal('0')) or Decimal('0')
            tabela_autos.cell(i, 5).text = f"{float(litros):,.2f}"
            tabela_autos.cell(i, 6).text = fmt_money(s.get("resto", Decimal('0')))

            # right align numeric cols
            for col_idx in range(1, len(cols)):
                try:
                    tabela_autos.cell(i, col_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                except Exception:
                    pass
    else:
        doc.add_paragraph("Nenhum registo por autocarro encontrado para o período.")

    doc.add_paragraph().add_run().add_break()

    # --- ASSINATURA --- (nome fixo e "cravado" no documento)
    assinatura_para = doc.add_paragraph()
    run = assinatura_para.add_run("Assinatura: ")
    run.bold = True
    nome_run = assinatura_para.add_run("KIANGEBENI KALEBA MATIAS")
    nome_run.bold = True

    doc.add_paragraph().add_run().add_break()

    # --- FINALIZAÇÃO E DOWNLOAD DO RELATÓRIO ---
    from io import BytesIO
    from django.http import HttpResponse

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Relatorio_Mensal_{mes_param}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    referer = request.META.get('HTTP_REFERER', '/')
    response['X-Redirect-After-Download'] = referer
    response['X-File-Generated'] = filename

    return response



# === ESTATÍSTICAS POR AUTOCARRO === #
@login_required
@acesso_restrito(['admin', 'gestor'])
def resumo_sector(request, slug):
    sector_obj = get_object_or_404(Sector, slug=slug)
    nivel = request.user.nivel_acesso.lower()

    if nivel == 'gestor' and sector_obj.gestor_id != request.user.id:
        return redirect('acesso_negado')
    elif nivel == 'associado' and not sector_obj.associados.filter(pk=request.user.pk).exists():
        return redirect('acesso_negado')
    elif nivel not in ['admin', 'gestor', 'associado']:
        return redirect('acesso_negado')

    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")

    # 🔹 Registos do setor
    registos = RegistoDiario.objects.filter(
        autocarro__sector=sector_obj
    ).select_related("autocarro")

    if data_inicio:
        registos = registos.filter(data__gte=parse_date(data_inicio))
    if data_fim:
        registos = registos.filter(data__lte=parse_date(data_fim))

    # 🔹 Entradas e saídas
    total_entradas = registos.aggregate(
        total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
    )["total"] or Decimal('0')

    total_saidas = registos.aggregate(
        total=Sum(F("alimentacao") + F("parqueamento") + F("outros"), output_field=DecimalField())
    )["total"] or Decimal('0')

    total_taxas = registos.aggregate(
        total=Sum("taxa", output_field=DecimalField())
    )["total"] or Decimal('0')

    total_km = registos.aggregate(Sum("km_percorridos"))["km_percorridos__sum"] or 0
    total_passageiros = registos.aggregate(Sum("numero_passageiros"))["numero_passageiros__sum"] or 0
    total_viagens = registos.aggregate(Sum("numero_viagens"))["numero_viagens__sum"] or 0

    # 🔹 Combustível
    combustivel_qs = DespesaCombustivel.objects.filter(autocarro__sector=sector_obj)
    if data_inicio:
        combustivel_qs = combustivel_qs.filter(data__gte=parse_date(data_inicio))
    if data_fim:
        combustivel_qs = combustivel_qs.filter(data__lte=parse_date(data_fim))

    comb_totais = combustivel_qs.aggregate(
        total_valor=Sum('valor', output_field=DecimalField()),
        total_litros=Sum('valor_litros', output_field=DecimalField()),
        total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
        total_lavagem=Sum('lavagem', output_field=DecimalField()),
    )

    total_combustivel_valor = comb_totais.get('total_valor') or Decimal('0')
    total_combustivel_sobragem = comb_totais.get('total_sobragem') or Decimal('0')
    total_combustivel_lavagem = comb_totais.get('total_lavagem') or Decimal('0')
    total_combustivel_geral = total_combustivel_valor + total_combustivel_sobragem + total_combustivel_lavagem

    # 🔹 Despesas variáveis
    despesas_qs = Despesa.objects.filter(sector=sector_obj)
    if data_inicio:
        despesas_qs = despesas_qs.filter(data__gte=parse_date(data_inicio))
    if data_fim:
        despesas_qs = despesas_qs.filter(data__lte=parse_date(data_fim))
    total_despesas_sector = despesas_qs.aggregate(total=Sum('valor', output_field=DecimalField())).get('total') or Decimal('0')

    # 🔹 Despesas fixas (também sujeitas a intervalo de datas)
    despesas_fixas = DespesaFixa.objects.filter(sector=sector_obj, ativo=True)

    if data_inicio:
        despesas_fixas = despesas_fixas.filter(data_inicio__gte=parse_date(data_inicio))
    if data_fim:
        despesas_fixas = despesas_fixas.filter(data_inicio__lte=parse_date(data_fim))

    total_despesas_fixas = despesas_fixas.aggregate(
        total=Sum('valor', output_field=DecimalField())
    ).get('total') or Decimal('0')

    # 🔹 Despesa geral (RelatorioSector) com filtro
    relatorios_qs = RelatorioSector.objects.filter(sector=sector_obj)
    if data_inicio:
        relatorios_qs = relatorios_qs.filter(data__gte=parse_date(data_inicio))
    if data_fim:
        relatorios_qs = relatorios_qs.filter(data__lte=parse_date(data_fim))
    despesa_geral_total = relatorios_qs.aggregate(total=Sum('despesa_geral', output_field=DecimalField())).get('total') or Decimal('0')

    # 🔹 Soma total de saídas
    total_saidas_final = (
        total_saidas
        + total_combustivel_geral
        + total_despesas_sector
        + total_despesas_fixas
        + despesa_geral_total
        + total_taxas
    )

    resto = total_entradas - total_saidas_final

    # 🔹 Estatísticas por autocarro
    autocarros_stats = []
    for autocarro in Autocarro.objects.filter(sector=sector_obj):
        registos_auto = registos.filter(autocarro=autocarro)
        entradas_auto = registos_auto.aggregate(
            total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
        )["total"] or Decimal('0')
        saidas_auto = registos_auto.aggregate(
            total=Sum(F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"), output_field=DecimalField())
        )["total"] or Decimal('0')

        comb_auto = combustivel_qs.filter(autocarro=autocarro).aggregate(
            total_valor=Sum('valor', output_field=DecimalField()),
            total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
            total_lavagem=Sum('lavagem', output_field=DecimalField()),
        )

        total_comb_auto = (
            (comb_auto.get('total_valor') or Decimal('0')) +
            (comb_auto.get('total_sobragem') or Decimal('0')) +
            (comb_auto.get('total_lavagem') or Decimal('0'))
        )

        saidas_auto += total_comb_auto
        resto_auto = entradas_auto - saidas_auto

        autocarros_stats.append({
            "autocarro": autocarro,
            "total_entradas": entradas_auto,
            "total_saidas": saidas_auto,
            "resto": resto_auto,
        })

    # 🔹 Formatação
    def fmt(valor):
        return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # 🔹 Mensagem WhatsApp formatada
    try:
        periodo = f"{data_inicio or '—'} até {data_fim or '—'}"
        msg_lines = [
            f"📊 *Resumo do Sector:* {sector_obj.nome}",
            f"📅 *Período:* {periodo}",
            "",
            f"💵 *Totais Gerais:*",
            f"• Entradas: {fmt(total_entradas)} Kz",
            f"• Saídas (registos): {fmt(total_saidas)} Kz",
            f"• Combustível: {fmt(total_combustivel_geral)} Kz",
            f"• Despesas Variáveis: {fmt(total_despesas_sector)} Kz",
            f"• Despesas Fixas: {fmt(total_despesas_fixas)} Kz",
            f"• Despesa Geral (Relatórios): {fmt(despesa_geral_total)} Kz",
            "",
            f"💰 *Total de Saídas:* {fmt(total_saidas_final)} Kz",
            f"💸 *Resto:* {fmt(resto)} Kz",
            "",
            "🚌 *Resumo por Autocarro:*",
        ]
        for s in autocarros_stats:
            msg_lines.append(
                f"• {s['autocarro'].numero}: Entradas {fmt(s['total_entradas'])} | "
                f"Saídas {fmt(s['total_saidas'])} | Resto {fmt(s['resto'])}"
            )
        msg_lines.append("")
        msg_lines.append(f"📅 *Gerado em:* {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        whatsapp_text = "\n".join(msg_lines)
        whatsapp_link = "https://api.whatsapp.com/send?text=" + quote_plus(whatsapp_text)
    except Exception:
        whatsapp_text = "Resumo do sector não disponível."
        whatsapp_link = "https://api.whatsapp.com/send?text=" + quote_plus(whatsapp_text)

    # 🔹 Contexto
    context = {
        "sector": sector_obj,
        "autocarros_stats": autocarros_stats,
        "total_entradas": total_entradas,
        "total_saidas": total_saidas_final,
        "total_km": total_km,
        "total_passageiros": total_passageiros,
        "total_viagens": total_viagens,
        "resto": resto,
        "total_combustivel_valor": total_combustivel_valor,
        "total_combustivel_sobragem": total_combustivel_sobragem,
        "total_combustivel_lavagem": total_combustivel_lavagem,
        "total_despesas_sector": total_despesas_sector,
        "total_despesas_fixas": total_despesas_fixas,
        "despesa_geral_total": despesa_geral_total,
        "total_taxas": total_taxas,
        "whatsapp_message": whatsapp_text,
        "whatsapp_link": whatsapp_link,
    }
    return render(request, "autocarros/resumo_sector.html", context)


# === DETALHE POR AUTOCARRO === #
@login_required
@acesso_restrito(['admin'])
def detalhe_autocarro(request, autocarro_id):
    autocarro = get_object_or_404(Autocarro, id=autocarro_id)
    registos_local = RegistoDiario.objects.filter(autocarro=autocarro)

    entradas = registos_local.aggregate(
        total=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"), output_field=DecimalField())
    )["total"] or 0

    saidas = registos_local.aggregate(
        total=Sum(F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"), output_field=DecimalField())
    )["total"] or 0

    km = registos_local.aggregate(Sum('km_percorridos'))['km_percorridos__sum'] or 0
    passageiros = registos_local.aggregate(Sum('numero_passageiros'))['numero_passageiros__sum'] or 0
    viagens = registos_local.aggregate(Sum('numero_viagens'))['numero_viagens__sum'] or 0
    # Agregar despesas de combustível por autocarro e por data
    combustiveis_qs = DespesaCombustivel.objects.filter(autocarro=autocarro, data__in=registos_local.values_list('data', flat=True))
    from collections import defaultdict
    comb_map = defaultdict(lambda: {
        'total_valor': Decimal('0'),
        'total_valor_litros': Decimal('0'),
        'total_sobragem': Decimal('0'),
        'total_lavagem': Decimal('0'),
    })
    total_combustivel_valor = Decimal('0')
    total_combustivel_litros = Decimal('0')
    total_combustivel_sobragem = Decimal('0')
    total_combustivel_lavagem = Decimal('0')

    for c in combustiveis_qs:
        key = c.data.isoformat()
        comb_map[key]['total_valor'] += c.valor or Decimal('0')
        comb_map[key]['total_valor_litros'] += c.valor_litros or Decimal('0')
        comb_map[key]['total_sobragem'] += c.sobragem_filtros or Decimal('0')
        comb_map[key]['total_lavagem'] += c.lavagem or Decimal('0')
        total_combustivel_valor += c.valor or Decimal('0')
        total_combustivel_litros += c.valor_litros or Decimal('0')
        total_combustivel_sobragem += c.sobragem_filtros or Decimal('0')
        total_combustivel_lavagem += c.lavagem or Decimal('0')

    # Anexar valores agregados a cada registo
    for r in registos_local:
        key = r.data.isoformat()
        agg = comb_map.get(key, {})
        r.combustivel_total = agg.get('total_valor', Decimal('0'))
        r.combustivel_valor_litros = agg.get('total_valor_litros', Decimal('0'))
        r.combustivel_sobragem = agg.get('total_sobragem', Decimal('0'))
        r.combustivel_lavagem = agg.get('total_lavagem', Decimal('0'))
        # saídas e saldo que incluem combustível
        try:
            r.saidas_total_incl_combustivel = r.saidas_total() + r.combustivel_total
        except Exception:
            r.saidas_total_incl_combustivel = r.saidas_total()
        try:
            r.saldo_liquido_incl_combustivel = r.entradas_total() - r.saidas_total_incl_combustivel
        except Exception:
            r.saldo_liquido_incl_combustivel = r.saldo_liquido()
        try:
            if r.combustivel_valor_litros and r.combustivel_valor_litros != Decimal('0'):
                r.preco_litro = r.combustivel_total / r.combustivel_valor_litros
            else:
                r.preco_litro = None
        except Exception:
            r.preco_litro = None

    resto = entradas - (saidas + total_combustivel_valor)

    # incluir sobragem/lavagem também no total de saídas apresentado
    total_saidas_incl_combustivel = saidas + total_combustivel_valor + total_combustivel_sobragem + total_combustivel_lavagem
    contexto = {
        'autocarro': autocarro,
        'total_entradas': entradas,
        'total_saidas': total_saidas_incl_combustivel,
        'total_km': km,
        'total_passageiros': passageiros,
        'total_viagens': viagens,
        'resto': resto,
        'total_combustivel_valor': total_combustivel_valor,
        'total_combustivel_litros': total_combustivel_litros,
        'total_combustivel_sobragem': total_combustivel_sobragem,
        'total_combustivel_lavagem': total_combustivel_lavagem,
        'registos': registos_local.order_by('-data')[:10],
    }
    return render(request, 'autocarros/detalhe_autocarro.html', contexto)


# === LISTAR REGISTROS DIÁRIOS COM FILTROS E AGRUPAMENTOS === #
from django.utils import timezone
from datetime import date
from decimal import Decimal
from django.db.models import Q, Sum
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib.auth.decorators import login_required
from urllib.parse import quote_plus
from .models import RegistoDiario, DespesaCombustivel, Sector


@login_required
def listar_registros(request):
    hoje = timezone.now().date()

    # Obter parâmetros de filtro
    sector_id = request.GET.get('sector', '').strip()
    data_inicio = request.GET.get('data_inicio', hoje.isoformat())
    data_fim = request.GET.get('data_fim', hoje.isoformat())

    nivel = request.user.nivel_acesso.lower()

    # 🔹 Determinar o setor (se foi selecionado)
    sector_obj = None
    if sector_id:
        sector_obj = get_object_or_404(Sector, id=sector_id)

    # 🔒 ---- Validação de acesso ----
    if nivel == 'gestor':
        if sector_obj:
            if sector_obj.gestor_id != request.user.id:
                return redirect('acesso_negado')
        else:
            sectores_permitidos = Sector.objects.filter(gestor=request.user)

    elif nivel == 'associado':
        if sector_obj:
            if not sector_obj.associados.filter(pk=request.user.pk).exists():
                return redirect('acesso_negado')
        else:
            sectores_permitidos = Sector.objects.filter(associados=request.user)

    elif nivel in ['admin', 'superuser']:
        sectores_permitidos = Sector.objects.all()

    else:
        return redirect('acesso_negado')

    # 🔹 ---- Consulta segura ----
    registros = RegistoDiario.objects.select_related('autocarro__sector')

    if sector_obj:
        registros = registros.filter(autocarro__sector=sector_obj)
    else:
        registros = registros.filter(autocarro__sector__in=sectores_permitidos)

    if data_inicio:
        registros = registros.filter(data__gte=data_inicio)
    if data_fim:
        registros = registros.filter(data__lte=data_fim)

    # 🔹 Agregar despesas de combustível
    combustivel_map = {}
    if registros.exists():
        autocarro_ids = set(registros.values_list('autocarro_id', flat=True))
        datas = set(registros.values_list('data', flat=True))
        combustiveis = DespesaCombustivel.objects.filter(
            autocarro_id__in=autocarro_ids,
            data__in=datas
        )

        from collections import defaultdict
        agg = defaultdict(lambda: {
            'total_valor': Decimal('0'),
            'total_valor_litros': Decimal('0'),
            'total_sobragem_filtros': Decimal('0'),
            'total_lavagem': Decimal('0'),
            'comprovativos': []
        })

        for c in combustiveis:
            key = f"{c.autocarro_id}_{c.data.isoformat()}"
            agg[key]['total_valor'] += c.valor or Decimal('0')
            agg[key]['total_valor_litros'] += c.valor_litros or Decimal('0')
            agg[key]['total_sobragem_filtros'] += c.sobragem_filtros or Decimal('0')
            agg[key]['total_lavagem'] += c.lavagem or Decimal('0')
            if c.comprovativo:
                url = c.comprovativo.url if hasattr(c.comprovativo, 'url') else str(c.comprovativo)
                agg[key]['comprovativos'].append(url)

        combustivel_map = dict(agg)

    # 🔹 Ordenar e agrupar
    registros = registros.order_by('-data', 'autocarro__sector__nome', 'autocarro__numero')
    registros_agrupados = {}

    for registro in registros:
        chave = f"{registro.data.isoformat()}_{registro.autocarro.sector.id}"
        
        if chave not in registros_agrupados:
            registros_agrupados[chave] = {
                'data': registro.data,
                'sector': registro.autocarro.sector,
                'registos': [],
                'total_entradas': Decimal('0'),
                'total_saidas': Decimal('0'),
                'total_saldo': Decimal('0'),
                'despesa_geral': Decimal('0'),
                'alimentacao_estaleiro': Decimal('0'),
            }

            # 🔹 Buscar o relatório apenas uma vez por grupo (otimizado)
            relatorio_sector = (
                RelatorioSector.objects
                .filter(sector=registro.autocarro.sector, data=registro.data)
                .only('despesa_geral', 'alimentacao_estaleiro')  # otimiza a query
                .first()
            )
            if relatorio_sector:
                registros_agrupados[chave]['despesa_geral'] = relatorio_sector.despesa_geral or Decimal('0')
                registros_agrupados[chave]['alimentacao_estaleiro'] = relatorio_sector.alimentacao_estaleiro or Decimal('0')
        # 🔹 Processamento do registro (combustível, entradas, saídas, etc)
        key = f"{registro.autocarro_id}_{registro.data.isoformat()}"
        comb = combustivel_map.get(key)

        # Combustível agregado
        if comb:
            registro.combustivel_total = comb.get('total_valor', Decimal('0'))
            registro.combustivel_valor_litros = comb.get('total_valor_litros', Decimal('0'))
            registro.combustivel_sobragem = comb.get('total_sobragem_filtros', Decimal('0'))
            registro.combustivel_lavagem = comb.get('total_lavagem', Decimal('0'))
            registro.comprovativos_combustivel = comb.get('comprovativos', [])
        else:
            registro.combustivel_total = Decimal('0')
            registro.combustivel_valor_litros = Decimal('0')
            registro.combustivel_sobragem = Decimal('0')
            registro.combustivel_lavagem = Decimal('0')
            registro.comprovativos_combustivel = []

        # Totais com combustível
        try:
            registro.saidas_total_incl_combustivel = (
                registro.saidas_total()
                + registro.combustivel_total
                + registro.combustivel_sobragem
                + registro.combustivel_lavagem
            )
        except Exception:
            registro.saidas_total_incl_combustivel = registro.saidas_total()

        try:
            registro.saldo_liquido_incl_combustivel = (
                registro.entradas_total() - registro.saidas_total_incl_combustivel
            )
        except Exception:
            registro.saldo_liquido_incl_combustivel = registro.saldo_liquido()

        # Preço por litro
        try:
            if registro.combustivel_valor_litros and registro.combustivel_valor_litros != Decimal('0'):
                registro.preco_litro = registro.combustivel_total / registro.combustivel_valor_litros
            else:
                registro.preco_litro = None
        except Exception:
            registro.preco_litro = None

        registros_agrupados[chave]['registos'].append(registro)
        registros_agrupados[chave]['total_entradas'] += registro.entradas_total()
        registros_agrupados[chave]['total_saidas'] += registro.saidas_total_incl_combustivel

    # ✅ Só agora calcula o saldo final de cada grupo (com despesa geral subtraída 1x)
    for grupo in registros_agrupados.values():
        # 🔹 Somar despesa_geral nas saídas totais
        grupo['total_saidas'] += grupo.get('despesa_geral', Decimal('0'))
        grupo['total_saidas'] += grupo.get('alimentacao_estaleiro', Decimal('0'))

        # 🔹 Agora calcula o saldo líquido
        grupo['total_saldo'] = (
            grupo['total_entradas'] - grupo['total_saidas']
        )

    # 🔹 Totais gerais
    total_entradas = sum(g['total_entradas'] for g in registros_agrupados.values())
    total_saidas = sum(g['total_saidas'] for g in registros_agrupados.values())
    total_saldo = sum(g['total_saldo'] for g in registros_agrupados.values())
    total_combustivel = sum(v.get('total_valor', Decimal('0')) for v in combustivel_map.values())
    total_despesas_gerais = sum(g.get('despesa_geral', Decimal('0')) for g in registros_agrupados.values())

    totais = {
        'total_entradas': total_entradas,
        'total_saidas': total_saidas,
        'total_saldo': total_saldo,
        'total_autocarros': registros.count(),
        'total_combustivel': total_combustivel,
        'total_despesas_gerais': total_despesas_gerais,
    }


    # 🔹 Gerar mensagens de WhatsApp
    def fmt_money(valor):
        try:
            d = Decimal(valor)
        except Exception:
            return "0,00"
        sign = '-' if d < 0 else ''
        d = abs(d).quantize(Decimal('0.01'))
        s = f"{d:.2f}"
        integer, frac = s.split('.')
        integer = '{:,}'.format(int(integer)).replace(',', '.')
        return f"{sign}{integer},{frac}"

    for g in registros_agrupados.values():
        data_str = g['data'].strftime('%d/%m/%Y')
        sector_name = g['sector'].nome if g.get('sector') else 'Geral'
        descricao = getattr(g['registos'][0].relatorio, 'descricao', '-') if g['registos'] else '-'

        parts = [
            f"Saudações!",
            f"",
            f"📊 RELATÓRIO DIÁRIO DE AUTOCARROS",
            f"📅 DATA: {data_str}",
            f"🏢 REGIÃO: {sector_name}",
            "",
            f"📝 DESCRIÇÃO: {descricao}",
        ]

        for reg in g['registos']:
            parts.append("")
            parts.append("__________________________________________")
            parts.append("")
            parts.append(f"🚌 Autocarro: {reg.autocarro.numero} - {reg.autocarro.modelo}")
            parts.append(f"👨‍✈️ Motorista: {reg.motorista or 'N/A'}")
            parts.append(f"👨‍💼 Cobrador Principal: {reg.cobrador_principal or 'N/A'}")
            parts.append(f"👨‍💼 Cobrador Auxiliar: {reg.cobrador_auxiliar or 'N/A'}")
            parts.append("")
            parts.append("✅ Entradas")
            parts.append(f"Normal: {fmt_money(getattr(reg, 'normal', 0))}kz")
            parts.append(f"Alunos: {fmt_money(getattr(reg, 'alunos', 0))}kz")
            parts.append(f"Luvu: {fmt_money(getattr(reg, 'luvu', 0))}kz")
            parts.append(f"Frete: {fmt_money(getattr(reg, 'frete', 0))}kz")
            parts.append(f"➡️ Total Entradas: {fmt_money(reg.entradas_total())}kz")
            parts.append("")
            parts.append("❌ Saídas")
            parts.append(f"Alimentação: {fmt_money(getattr(reg, 'alimentacao', 0))}kz")
            parts.append(f"Parqueamento: {fmt_money(getattr(reg, 'parqueamento', 0))}kz")
            parts.append(f"Taxa: {fmt_money(getattr(reg, 'taxa', 0))}kz")
            parts.append(f"Outros: {fmt_money(getattr(reg, 'outros', 0))}kz")
            parts.append(f"Taxi: {fmt_money(getattr(reg, 'taxi', 0))}kz")
            parts.append(f"Combustível: {fmt_money(reg.combustivel_total)}kz")
            parts.append(f"Sobragem/Filtros: {fmt_money(reg.combustivel_sobragem)}kz")
            parts.append(f"Lavagem: {fmt_money(reg.combustivel_lavagem)}kz")
            parts.append(f"➡️ Total Saídas: {fmt_money(reg.saidas_total_incl_combustivel)}kz")
            parts.append("")
            parts.append("📊 Outros Dados")
            parts.append(f"Kms: {getattr(reg, 'km_percorridos', 0)}")
            parts.append(f"Passageiros: {getattr(reg, 'numero_passageiros', 0)}")
            parts.append(f"Viagens: {getattr(reg, 'numero_viagens', 0)}")
            parts.append(f"💰 Saldo Liquído: {fmt_money(reg.saldo_liquido_incl_combustivel)}kz")

        parts.append("")
        parts.append("____________________________")
        parts.append("")
        parts.append("📊 Resumo Geral")
        parts.append(f"✅ Entrada Geral: {fmt_money(total_entradas)}kz")
        parts.append(f"❌Despesa Feita Na Produção: {fmt_money(g.get('despesa_geral', 0))}kz")
        parts.append(f"❌Alimentação Estaleiro: {fmt_money(g.get('alimentacao_estaleiro', 0))}kz")
        parts.append(f"❌ Saída Geral: {fmt_money(total_saidas)}kz")
        parts.append(f"💰 Liquído Geral: {fmt_money(total_saldo)}kz")
        parts.append("")
        parts.append("Suporte técnico: @kiangebenimatias4@gmail.com, +244 944 790 744 (WhatsApp)")

        message = '\n'.join(parts)
        g['whatsapp_link'] = f"https://wa.me/?text={quote_plus(message)}"

    # 🔹 Contexto final
    sectores = Sector.objects.all()
    context = {
        'registros_agrupados': list(registros_agrupados.values()),
        'sectores': sectores,
        'sector_id': sector_id,
        'data_inicio': data_inicio,
        'data_fim': data_fim,
        'totais': totais,
        'hoje': hoje,
        'combustivel_map': combustivel_map,
    }

    return render(request, 'autocarros/listar_registros.html', context)


@login_required
def deletar_registros_sector_data(request, sector_id, data):
    sector = get_object_or_404(Sector, pk=sector_id)
    data_obj = parse_date(data)
    if not data_obj:
        messages.error(request, '❌ Data inválida.')
        return redirect('listar_registros')
    if request.method == 'POST':
        try:
            # Excluir todos os registros do setor e data
            RegistoDiario.objects.filter(autocarro__sector=sector, data=data_obj).delete()
            # Excluir o relatório do setor e data
            RelatorioSector.objects.filter(sector=sector, data=data_obj).delete()
            messages.success(request, f'✅ Todos os registros e o relatório do setor {sector.nome} em {data_obj.strftime("%d/%m/%Y")} foram eliminados!')
            return redirect('listar_registros')
        except Exception as e:
            messages.error(request, f'❌ Erro ao eliminar registros: {str(e)}')
    return render(request, 'autocarros/confirmar_deletar_registros_sector.html', {'sector': sector, 'data': data_obj})


@login_required
def deletar_registro(request, pk):
    registro = get_object_or_404(RegistoDiario, pk=pk)
    if request.method == 'POST':
        try:
            registro.delete()
            messages.success(request, '✅ Registro eliminado com sucesso!')
            return redirect('listar_registros')
        except Exception as e:
            messages.error(request, f'❌ Erro ao eliminar registro: {str(e)}')
    return render(request, 'autocarros/confirmar_deletar_registro.html', {'registro': registro})


@login_required
def concluir_relatorio(request, pk):
    """Marca o relatório como concluído"""
    relatorio = get_object_or_404(RelatorioSector, pk=pk)
    
    if request.method == 'POST':
        try:
            # 🔹 Aqui você pode adicionar lógica adicional se necessário
            messages.success(request, f"✅ Relatório de {relatorio.sector.nome} concluído com sucesso!")
        except Exception as e:
            messages.error(request, f"❌ Erro ao concluir relatório: {str(e)}")
    
    return redirect('listar_registros')


@login_required
@acesso_restrito(['admin'])
def validar_relatorio(request, pk):
    """Marca o relatório como validado pelo supervisor — só admin pode validar."""
    relatorio = get_object_or_404(RelatorioSector, pk=pk)

    # Verificação redundante/explicita para evitar inconsistências no campo nivel_acesso
    nivel = getattr(request.user, "nivel_acesso", "") or ""
    if not (hasattr(request.user, "is_admin") and request.user.is_admin()) and nivel.lower() != "admin":
        messages.error(request, "❌ Acesso negado. Apenas administradores podem validar relatórios.")
        return redirect('acesso_negado')

    if request.method == 'POST':
        try:
            # marcar como validado (implemente a lógica real aqui, ex: relatorio.validado = True; relatorio.save())
            # Exemplo genérico:
            if hasattr(relatorio, "validado"):
                relatorio.validado = True
                relatorio.save()
            messages.success(request, f"✅ Relatório de {relatorio.sector.nome} validado com sucesso!")
        except Exception as e:
            messages.error(request, f"❌ Erro ao validar relatório: {str(e)}")

    return redirect('listar_registros')


@login_required
def relatorios_validados(request):
    # Obter parâmetros de filtro
    sector_id = request.GET.get('sector', '')
    data_inicio = request.GET.get('data_inicio', '')
    data_fim = request.GET.get('data_fim', '')
    
    # Data padrão: hoje
    data_hoje = timezone.now().date()
    
    # 🔹 CORRIGIR a query - agora trabalhamos diretamente com RegistoDiario
    registos_validados = RegistoDiario.objects.filter(validado=True)
    
    # Aplicar filtros
    if sector_id:
        registos_validados = registos_validados.filter(autocarro__sector_id=sector_id)
    
    if data_inicio:
        registos_validados = registos_validados.filter(data__gte=data_inicio)
    
    if data_fim:
        registos_validados = registos_validados.filter(data__lte=data_fim)
    
    # Se não há filtros de data, mostrar apenas dados do dia atual
    if not data_inicio and not data_fim:
        registos_validados = registos_validados.filter(data=data_hoje)
        data_inicio = data_hoje.isoformat()
        data_fim = data_hoje.isoformat()
    
    # Agregar despesas de combustível para os registos validados (por autocarro+data)
    combustivel_map_validados = {}
    if registos_validados.exists():
        autocarro_ids = set(registos_validados.values_list('autocarro_id', flat=True))
        datas = set(registos_validados.values_list('data', flat=True))
        combustiveis_val = DespesaCombustivel.objects.filter(autocarro_id__in=autocarro_ids, data__in=datas)
        from collections import defaultdict
        agg_val = defaultdict(lambda: {
            'total_valor': Decimal('0'),
            'total_valor_litros': Decimal('0'),
            'total_sobragem': Decimal('0'),
            'total_lavagem': Decimal('0'),
        })
        for c in combustiveis_val:
            key = f"{c.autocarro_id}_{c.data.isoformat()}"
            agg_val[key]['total_valor'] += c.valor or Decimal('0')
            agg_val[key]['total_valor_litros'] += c.valor_litros or Decimal('0')
            agg_val[key]['total_sobragem'] += c.sobragem_filtros or Decimal('0')
            agg_val[key]['total_lavagem'] += c.lavagem or Decimal('0')
        for k, v in agg_val.items():
            combustivel_map_validados[k] = v

    # Agrupar registos por data e sector para exibição e anexar combustíveis
    registos_por_data_sector = {}
    processed_registos = []
    for registro in registos_validados.select_related('autocarro__sector'):
        # anexar agregados de combustível ao registro para uso nos totais
        key = f"{registro.autocarro_id}_{registro.data.isoformat()}"
        comb = combustivel_map_validados.get(key, {})
        registro.combustivel_total = comb.get('total_valor', Decimal('0'))
        registro.combustivel_valor_litros = comb.get('total_valor_litros', Decimal('0'))
        registro.combustivel_sobragem = comb.get('total_sobragem', Decimal('0'))
        registro.combustivel_lavagem = comb.get('total_lavagem', Decimal('0'))
        # saídas e saldo que incluem combustível
        try:
            registro.saidas_total_incl_combustivel = registro.saidas_total() + registro.combustivel_total + registro.combustivel_sobragem + registro.combustivel_lavagem
        except Exception:
            registro.saidas_total_incl_combustivel = registro.saidas_total()
        try:
            registro.saldo_liquido_incl_combustivel = registro.entradas_total() - registro.saidas_total_incl_combustivel
        except Exception:
            registro.saldo_liquido_incl_combustivel = registro.saldo_liquido()

        chave = f"{registro.data}_{registro.autocarro.sector.id}"
        if chave not in registos_por_data_sector:
            registos_por_data_sector[chave] = {
                'data': registro.data,
                'sector': registro.autocarro.sector,
                'registos': [],
                'total_entradas': Decimal('0'),
                'total_saidas': Decimal('0'),
                'total_saldo': Decimal('0'),
                'total_combustivel': Decimal('0'),
            }

        # atualizar totais do grupo usando valores já anexados ao registro
        try:
            entradas_reg = registro.entradas_total()
        except Exception:
            entradas_reg = Decimal('0')
        # garantir Decimal
        try:
            entradas_reg = Decimal(entradas_reg)
        except Exception:
            entradas_reg = Decimal('0')

        try:
            saidas_reg = getattr(registro, 'saidas_total_incl_combustivel', registro.saidas_total())
        except Exception:
            saidas_reg = Decimal('0')
        try:
            saidas_reg = Decimal(saidas_reg)
        except Exception:
            saidas_reg = Decimal('0')

        try:
            saldo_reg = getattr(registro, 'saldo_liquido_incl_combustivel', registro.saldo_liquido())
        except Exception:
            saldo_reg = Decimal('0')
        try:
            saldo_reg = Decimal(saldo_reg)
        except Exception:
            saldo_reg = Decimal('0')

        combustivel_reg = getattr(registro, 'combustivel_total', Decimal('0'))
        try:
            combustivel_reg = Decimal(combustivel_reg)
        except Exception:
            combustivel_reg = Decimal('0')

        registos_por_data_sector[chave]['registos'].append(registro)
        registos_por_data_sector[chave]['total_entradas'] += entradas_reg
        registos_por_data_sector[chave]['total_saidas'] += saidas_reg
        registos_por_data_sector[chave]['total_saldo'] += saldo_reg
        registos_por_data_sector[chave]['total_combustivel'] += combustivel_reg
        processed_registos.append(registro)
    
    # Calcular totais
    # Calcular totais a partir dos registros já processados (com combustíveis anexados)
    totais = {
        'total_entradas': sum(getattr(reg, 'entradas_total')() if callable(getattr(reg, 'entradas_total', None)) else Decimal('0') for reg in processed_registos) if processed_registos else Decimal('0'),
        'total_saidas': sum(getattr(reg, 'saidas_total_incl_combustivel', reg.saidas_total()) for reg in processed_registos) if processed_registos else Decimal('0'),
        'total_saldo': sum(getattr(reg, 'saldo_liquido_incl_combustivel', reg.saldo_liquido()) for reg in processed_registos) if processed_registos else Decimal('0'),
        'total_autocarros': len(processed_registos),
        'total_combustivel': sum(getattr(reg, 'combustivel_total', Decimal('0')) for reg in processed_registos) if processed_registos else Decimal('0'),
    }

    context = {
        'registos_agrupados': list(registos_por_data_sector.values()),
        'sectores': Sector.objects.all(),
        'sector_id': sector_id,
        'data_inicio': data_inicio,
        'data_fim': data_fim,
        'totais': totais,
        'data_hoje': data_hoje,
    }

    return render(request, 'autocarros/relatorios_validados.html', context)


@login_required
def adicionar_relatorio_sector(request):
    if request.method == 'POST':
        relatorio_form = RelatorioSectorForm(request.POST)
        
        # 🔹 USAR O FORMULÁRIO SIMPLIFICADO
        multi_file_form = MultiFileForm(request.POST, request.FILES)
        
        if relatorio_form.is_valid() and multi_file_form.is_valid():
            # 🔹 VERIFICAR SE JÁ EXISTE RELATÓRIO PARA ESTE SETOR NA DATA
            sector = relatorio_form.cleaned_data['sector']
            data = relatorio_form.cleaned_data['data']
            
            if RelatorioSector.objects.filter(sector=sector, data=data).exists():
                messages.error(request, f"❌ Já existe um relatório para o setor {sector.nome} na data {data}.")
                return render(request, 'autocarros/adicionar_relatorio_sector.html', {
                    'relatorio_form': relatorio_form,
                    'multi_file_form': multi_file_form
                })
            
            try:
                # Salvar o relatório
                relatorio = relatorio_form.save()
                
                # 🔹 SALVAR MÚLTIPLOS ARQUIVOS
                arquivos = request.FILES.getlist('arquivos')
                for arquivo in arquivos:
                    if arquivo:  # Verificar se o arquivo não está vazio
                        ComprovativoRelatorio.objects.create(
                            relatorio=relatorio,
                            arquivo=arquivo,
                            descricao=f"Comprovativo {arquivo.name}"
                        )
                
                # Criar registos para cada autocarro do sector
                autocarros = Autocarro.objects.filter(sector=relatorio.sector)
                registros_criados = []
                for autocarro in autocarros:
                    registro, criado = RegistoDiario.objects.get_or_create(
                        relatorio=relatorio,
                        autocarro=autocarro,
                        data=relatorio.data
                    )
                    if criado:
                        registros_criados.append(registro)

                messages.success(request, f"✅ Relatório para {relatorio.sector.nome} criado com {len(arquivos)} comprovativos!")
                # Redirecionar para o primeiro registro criado
                if registros_criados:
                    return redirect('editar_relatorio_sector', pk=registros_criados[0].pk)
                else:
                    return redirect('listar_registros')
                
            except Exception as e:
                messages.error(request, f"❌ Erro ao criar relatório: {str(e)}")
        else:
            messages.error(request, "❌ Erro no formulário. Verifique os dados.")
    else:
        relatorio_form = RelatorioSectorForm()
        multi_file_form = MultiFileForm()

    # Relatórios recentes para referência
    relatorios_recentes = RelatorioSector.objects.select_related('sector').order_by('-data')[:5]

    return render(request, 'autocarros/adicionar_relatorio_sector.html', {
        'relatorio_form': relatorio_form,
        'multi_file_form': multi_file_form,
        'relatorios_recentes': relatorios_recentes
    })


@login_required
def gerir_relatorio_sector(request, pk=None):
    relatorio = RelatorioSector.objects.filter(pk=pk).first()  # 🔹 Se for edição, busca o relatório existente

    if request.method == 'POST':
        relatorio_form = RelatorioSectorForm(request.POST, instance=relatorio)
        multi_file_form = MultiFileForm(request.POST, request.FILES)

        if relatorio_form.is_valid() and multi_file_form.is_valid():
            try:
                relatorio = relatorio_form.save()

                # 🔹 Atualizar ou criar comprovativos
                arquivos = request.FILES.getlist('arquivos')
                for arquivo in arquivos:
                    if arquivo:
                        ComprovativoRelatorio.objects.create(
                            relatorio=relatorio,
                            arquivo=arquivo,
                            descricao=f"Comprovativo {arquivo.name}"
                        )

                # 🔹 Criar registos (somente se for novo)
                if pk is None:
                    autocarros = Autocarro.objects.filter(sector=relatorio.sector)
                    for autocarro in autocarros:
                        RegistoDiario.objects.get_or_create(
                            relatorio=relatorio,
                            autocarro=autocarro,
                            data=relatorio.data
                        )

                if pk:
                    messages.success(request, f"✅ Relatório de {relatorio.sector.nome} atualizado com sucesso!")
                else:
                    messages.success(request, f"✅ Relatório de {relatorio.sector.nome} criado com sucesso!")

                return redirect('gerir_relatorio_sector', pk=relatorio.pk)

            except Exception as e:
                messages.error(request, f"❌ Erro ao salvar relatório: {e}")
        else:
            messages.error(request, "❌ Corrija os erros do formulário.")

    else:
        relatorio_form = RelatorioSectorForm(instance=relatorio)
        multi_file_form = MultiFileForm()

    relatorios_recentes = RelatorioSector.objects.select_related('sector').order_by('-data')[:5]

    return render(request, 'autocarros/adicionar_relatorio_sector.html', {
        'relatorio_form': relatorio_form,
        'multi_file_form': multi_file_form,
        'relatorios_recentes': relatorios_recentes,
        'relatorio': relatorio,
        'modo_edicao': bool(relatorio),
    })


@login_required
@acesso_restrito(['admin', 'gestor'])
def editar_registros(request, sector_id, data):
    return editar_relatorio_sector(request, sector_id=sector_id, data=data)


@login_required
@acesso_restrito(['admin', 'gestor'])
def editar_relatorio_sector(request, pk=None, sector_id=None, data=None):
    """
    Permite editar tanto o relatório do setor quanto os registros individuais.
    Pode ser acessada:
      - Por PK de um RegistoDiario (editar_relatorio_sector)
      - Por sector_id + data (editar_registros)
    """
    # 🔹 Determinar o sector e a data
    if pk:
        registro_base = get_object_or_404(RegistoDiario, pk=pk)
        sector = registro_base.autocarro.sector
        data = registro_base.data
    else:
        sector = get_object_or_404(Sector, pk=sector_id)
        data = parse_date(data)

    # 🔹 Buscar ou criar relatório do setor
    relatorio_sector, _ = RelatorioSector.objects.get_or_create(
        sector=sector,
        data=data,
        defaults={"despesa_geral": 0}
    )

    # 🔹 Garantir que todos os autocarros do setor tenham registro
    autocarros_do_sector = Autocarro.objects.filter(sector=sector)
    registros = RegistoDiario.objects.filter(
        autocarro__sector=sector,
        data=data
    ).select_related('autocarro')

    autocarros_com_registro = registros.values_list('autocarro_id', flat=True)
    for autocarro in autocarros_do_sector:
        if autocarro.id not in autocarros_com_registro:
            RegistoDiario.objects.create(autocarro=autocarro, data=data)

    # Recarregar registros
    registros = RegistoDiario.objects.filter(
        autocarro__sector=sector,
        data=data
    ).select_related('autocarro')

    # 🔹 POST - salvar alterações
    if request.method == "POST":
        relatorio_form = RelatorioSectorForm(request.POST, request.FILES, instance=relatorio_sector)

        if relatorio_form.is_valid():
            relatorio_form.save()

            # 🔹 Salvar novos comprovativos, se enviados
            novos_arquivos = request.FILES.getlist("novos_comprovativos")
            for arquivo in novos_arquivos:
                if arquivo:
                    ComprovativoRelatorio.objects.create(
                        relatorio=relatorio_sector,
                        arquivo=arquivo,
                        descricao=f"Comprovativo {arquivo.name}"
                    )
        else:
            messages.warning(request, "⚠️ Verifique o campo 'Despesa geral do setor'.")

        # 🔹 Salvar cada registo de autocarro
        for registro in registros:
            form = RegistoDiarioForm(
                request.POST,
                instance=registro,
                prefix=f"registro_{registro.id}"
            )

            if form.is_valid():
                try:
                    form.save()
                except Exception as e:
                    messages.error(
                        request,
                        f"Erro ao salvar registo do autocarro {registro.autocarro.numero}: {str(e)}"
                    )
            else:
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.warning(
                            request,
                            f"Autocarro {registro.autocarro.numero}, campo {field}: {error}"
                        )

        if not any(message.tags == 'error' for message in messages.get_messages(request)):
            messages.success(
                request,
                f"✅ Registos e relatório do setor {sector.nome} ({data}) atualizados com sucesso!"
            )
            return redirect("listar_registros")

    else:
        relatorio_form = RelatorioSectorForm(instance=relatorio_sector)

    # 🔹 Preparar contexto
    forms = []
    for registro in registros:
        forms.append({
            "form": RegistoDiarioForm(instance=registro, prefix=f"registro_{registro.id}"),
            "registro": registro,
        })

    context = {
        "sector": sector,
        "data": data,
        "relatorio": relatorio_sector,  # ✅ adicionado — essencial para o template
        "relatorio_form": relatorio_form,
        "forms": forms,
        "total_autocarros": autocarros_do_sector.count(),
        "total_registros": registros.count(),
    }

    return render(request, "autocarros/editar_relatorio_sector.html", context)


@login_required
def adicionar_comprovativos(request, pk):
    """Adicionar comprovativos a um relatório existente"""
    relatorio = get_object_or_404(RelatorioSector, pk=pk)
    
    if request.method == 'POST':
        arquivos = request.FILES.getlist('arquivos')
        descricao_geral = request.POST.get('descricao_geral', '')
        
        if arquivos:
            try:
                for arquivo in arquivos:
                    ComprovativoRelatorio.objects.create(
                        relatorio=relatorio,
                        arquivo=arquivo,
                        descricao=descricao_geral or f"Comprovativo {arquivo.name}"
                    )
                
                messages.success(request, f"✅ {len(arquivos)} comprovativo(s) adicionado(s) com sucesso!")
            except Exception as e:
                messages.error(request, f"❌ Erro ao adicionar comprovativos: {str(e)}")
        else:
            messages.warning(request, "⚠️ Nenhum arquivo selecionado.")
        
        return redirect('editar_relatorio_sector', pk=relatorio.pk)
    
    return redirect('editar_relatorio_sector', pk=relatorio.pk)


@login_required
def deletar_comprovativo(request, pk):
    """Deletar um comprovativo específico"""
    comprovativo = get_object_or_404(ComprovativoRelatorio, pk=pk)
    relatorio_pk = comprovativo.relatorio.pk
    
    if request.method == 'POST':
        try:
            comprovativo.delete()
            messages.success(request, "✅ Comprovativo excluído com sucesso!")
        except Exception as e:
            messages.error(request, f"❌ Erro ao excluir comprovativo: {str(e)}")
    
    return redirect('editar_relatorio_sector', pk=relatorio_pk)


@login_required
def deletar_relatorio_sector(request, pk):
    relatorio = get_object_or_404(RelatorioSector, pk=pk)

    if request.method == "POST":
        try:
            relatorio.delete()
            messages.success(request, "✅ Relatório do setor apagado com sucesso!")
            return redirect("listar_registros")
        except Exception as e:
            messages.error(request, f"❌ Erro ao apagar relatório: {str(e)}")

    return render(request, "autocarros/deletar_relatorio_sector.html", {
        "relatorio": relatorio
    })


# === Autocarros Views === #
@login_required
def listar_autocarros(request):
    autocarros = Autocarro.objects.all().order_by('numero')
    return render(request, 'autocarros/listar_autocarros.html', {'autocarros': autocarros})


@login_required
def alterar_status_autocarro(request, pk):
    autocarro = get_object_or_404(Autocarro, pk=pk)
    if request.method == "POST":
        novo_status = request.POST.get("status")
        if novo_status in dict(Autocarro._meta.get_field("status").choices):
            autocarro.status = novo_status
            autocarro.save()
            messages.success(request, f"✅ Status do autocarro {autocarro.numero} atualizado para {autocarro.get_status_display()}.")
        else:
            messages.error(request, "❌ Status inválido.")
    return redirect("listar_autocarros")


@login_required
def cadastrar_autocarro(request):
    if request.method == 'POST':
        form = AutocarroForm(request.POST)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, '✅ Autocarro cadastrado com sucesso!')
                return redirect('listar_autocarros')
            except Exception as e:
                messages.error(request, f'❌ Erro ao cadastrar autocarro: {str(e)}')
        else:
            messages.error(request, '❌ Erro no formulário. Verifique os dados.')
    else:
        form = AutocarroForm()
    return render(request, 'autocarros/cadastrar_autocarro.html', {'form': form})


@login_required
# Atualizar estado do autocarro
def atualizar_estado(request):
    if request.method == "POST":
        form = EstadoAutocarroForm(request.POST)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, "✅ Estado do autocarro atualizado com sucesso!")
                return redirect("listar_autocarros")
            except Exception as e:
                messages.error(request, f"❌ Erro ao atualizar estado: {str(e)}")
        else:
            messages.error(request, "❌ Erro no formulário. Verifique os dados.")
    else:
        form = EstadoAutocarroForm()
    return render(request, "autocarros/atualizar_estado.html", {"form": form})


@login_required
# editar autocarro
def editar_autocarro(request, pk):
    autocarro = get_object_or_404(Autocarro, pk=pk)
    if request.method == 'POST':
        form = AutocarroForm(request.POST, instance=autocarro)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, '✅ Autocarro atualizado com sucesso!')
               
                return redirect('listar_autocarros')
            except Exception as e:
                messages.error(request, f'❌ Erro ao atualizar autocarro: {str(e)}')
        else:
            messages.error(request, '❌ Erro no formulário. Verifique os dados.')
    else:
        form = AutocarroForm(instance=autocarro)
    return render(request, 'autocarros/editar_autocarro.html', {'form': form, 'autocarro': autocarro})


@login_required
# deletar autocarro
def deletar_autocarro(request, pk):
    autocarro = get_object_or_404(Autocarro, pk=pk)
    if request.method == 'POST':
        try:
            autocarro.delete()
            messages.success(request, '✅ Autocarro deletado com sucesso!')
            return redirect('listar_autocarros')
        except Exception as e:
            messages.error(request, f'❌ Erro ao deletar autocarro: {str(e)}')
    return render(request, 'autocarros/deletar_autocarro.html', {'autocarro': autocarro})


# ---------- Despesa Combustível Views ----------#
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import DespesaCombustivel, Autocarro


@login_required
# Listar despesas de combustível
def listar_combustivel(request):
    """
    Listagem de despesas de combustível (filtro por intervalo de datas e por autocarro).
    """
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")
    autocarro_id = request.GET.get("autocarro")

    qs = DespesaCombustivel.objects.select_related("autocarro", "autocarro__sector").all()
    if data_inicio:
        qs = qs.filter(data__gte=data_inicio)
    if data_fim:
        qs = qs.filter(data__lte=data_fim)
    if autocarro_id:
        qs = qs.filter(autocarro_id=autocarro_id)

    qs = qs.order_by("-data")

    autocarros = Autocarro.objects.all().order_by("numero")

    context = {
        "combustiveis": qs,
        "autocarros": autocarros,
        "data_inicio": data_inicio,
        "data_fim": data_fim,
        "autocarro_selected": autocarro_id,
    }
    return render(request, "despesas/listar_combustivel.html", context)


@login_required
# selecionar sector antes de adicionar despesas de combustível
def selecionar_sector_combustivel(request):
    if request.method == "POST":
        form = SelecionarSectorCombustivelForm(request.POST)
        if form.is_valid():
            sector = form.cleaned_data["sector"]
            return redirect("adicionar_combustivel", pk=sector.pk)
    else:
        form = SelecionarSectorCombustivelForm()

    return render(request, "despesas/selecionar_sector.html", {"form": form})


@login_required
# adicionar despesas de combustível para todos os autocarros de um sector
def adicionar_combustivel(request, pk):
    sector = get_object_or_404(Sector, pk=pk)
    autocarros = Autocarro.objects.filter(sector=sector).order_by("numero")

    CombustivelFormSet = modelformset_factory(
        DespesaCombustivel,
        form=DespesaCombustivelForm,
        extra=len(autocarros),
        can_delete=False
    )

    if request.method == "POST":
        formset = CombustivelFormSet(
            request.POST,
            request.FILES,
            queryset=DespesaCombustivel.objects.none()
        )
        if formset.is_valid():
            try:
                for form, autocarro in zip(formset.forms, autocarros):
                    if form.cleaned_data:
                        inst = form.save(commit=False)
                        inst.sector = sector
                        inst.autocarro = autocarro
                        if not inst.data:
                            inst.data = timezone.now().date()
                        inst.save()
                messages.success(request, "✅ Despesas de combustível adicionadas com sucesso!")
                return redirect("listar_combustivel")
            except Exception as e:
                messages.error(request, f"❌ Erro ao adicionar combustível: {str(e)}")
        else:
            messages.error(request, "❌ Erro no formulário. Verifique os dados.")
    else:
        initial_data = [{"sector": sector, "autocarro": a} for a in autocarros]
        formset = CombustivelFormSet(
            queryset=DespesaCombustivel.objects.none(),
            initial=initial_data
        )

    formset_autocarros = zip(formset.forms, autocarros)

    return render(request, "despesas/adicionar_combustivel.html", {
        "sector": sector,
        "formset": formset,
        "formset_autocarros": formset_autocarros,
    })


@login_required
# editar despesas de combustível
def editar_combustivel(request, pk):
    despesa = get_object_or_404(DespesaCombustivel, pk=pk)

    if request.method == "POST":
        form = DespesaCombustivelForm(request.POST, request.FILES, instance=despesa)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, "✅ Despesa de combustível atualizada com sucesso!")
                return redirect("listar_despesas")
            except Exception as e:
                messages.error(request, f"❌ Erro ao atualizar combustível: {str(e)}")
        else:
            messages.error(request, "❌ Erro no formulário. Verifique os dados.")
    else:
        form = DespesaCombustivelForm(instance=despesa)

    return render(request, "despesas/editar_combustivel.html", {
        "form": form,
        "despesa": despesa,
    })


@login_required
# deletar despesas de combustível
def deletar_combustivel(request, pk):
    despesa = get_object_or_404(DespesaCombustivel, pk=pk)

    if request.method == "POST":
        try:
            despesa.delete()
            messages.success(request, "✅ Despesa de combustível apagada com sucesso!")
            return redirect("listar_combustivel")
        except Exception as e:
            messages.error(request, f"❌ Erro ao apagar combustível: {str(e)}")

    return render(request, "despesas/deletar_combustivel.html", {
        "despesa": despesa,
    })


# ------------ Despesas normais ou variavéis ----------#
@login_required
# Despesas 'normais ou variáveis' (não combustível)
def listar_despesas(request):
    """
    Listar apenas despesas 'normais' (classificadas aqui como 'variáveis').
    Filtra por intervalo de datas se fornecido.
    """
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")

    despesas_qs = Despesa.objects.all()
    if data_inicio:
        despesas_qs = despesas_qs.filter(data__gte=data_inicio)
    if data_fim:
        despesas_qs = despesas_qs.filter(data__lte=data_fim)

    # Ordenar por data decrescente e empacotar para o template
    despesas_qs = despesas_qs.order_by('-data')
    despesas = [{"tipo": "variavel", "obj": d} for d in despesas_qs]

    return render(request, "despesas/listar_despesas.html", {"despesas": despesas, "data_inicio": data_inicio, "data_fim": data_fim})


@login_required
@acesso_restrito(['admin'])
# adicionar despesas 'normais ou variáveis'
def adicionar_despesa(request):
    if request.method == 'POST':
        form = DespesaForm(request.POST)
        multi = MultiFileForm(request.POST, request.FILES)

        if form.is_valid() and multi.is_valid():
            try:
                despesa = form.save()
                arquivos = request.FILES.getlist('arquivos')
                for arquivo in arquivos:
                    if arquivo:
                        Comprovativo.objects.create(despesa=despesa, arquivo=arquivo)
                messages.success(request, '✅ Despesa adicionada com sucesso!')
                return redirect('listar_despesas')
            except Exception as e:
                messages.error(request, f'❌ Erro ao adicionar despesa: {str(e)}')
        else:
            messages.error(request, '❌ Erro no formulário. Verifique os dados.')
    else:
        form = DespesaForm()
        multi = MultiFileForm()

    return render(request, 'despesas/adicionar_despesa.html', {'form': form, 'multi': multi})


@login_required
# editar despesas 'normais ou variáveis'
@acesso_restrito(['admin'])
def editar_despesa(request, pk):
    try:
        despesa = Despesa.objects.get(pk=pk)
    except Despesa.DoesNotExist:
        messages.error(request, "A despesa que tentou editar não existe ou foi removida.")
        return redirect('listar_despesas')
    
    if request.method == 'POST':
        form = DespesaForm(request.POST, instance=despesa)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, '✅ Despesa atualizada com sucesso!')
                return redirect('listar_despesas')
            except Exception as e:
                messages.error(request, f'❌ Erro ao atualizar despesa: {str(e)}')
        else:
            messages.error(request, '❌ Erro no formulário. Verifique os dados.')
    else:
        form = DespesaForm(instance=despesa)
    return render(request, 'despesas/editar_despesa.html', {'form': form, 'despesa': despesa})


@login_required
# deletar despesas 'normais ou variáveis'
@acesso_restrito(['admin'])
def deletar_despesa(request, pk):
    try:
        despesa = Despesa.objects.get(pk=pk)
    except Despesa.DoesNotExist:
        messages.error(request, "A despesa que tentou eliminar não existe ou já foi removida.")
        return redirect('listar_despesas')

    if request.method == 'POST':
        try:
            despesa.delete()
            messages.success(request, "✅ Despesa eliminada com sucesso!")
            return redirect('listar_despesas')
        except Exception as e:
            messages.error(request, f"❌ Erro ao eliminar a despesa: {str(e)}")

    return render(request, 'despesas/deletar_despesa.html', {'despesa': despesa})


# ---------- Despesas Fixas Views ----------#
@login_required
@acesso_restrito(['admin'])
def listar_despesas_fixas(request):
    """Listagem de despesas fixas (filtro por sector opcional)."""
    sector_id = request.GET.get('sector')
    qs = DespesaFixa.objects.select_related('sector','responsavel').all()
    if sector_id:
        qs = qs.filter(sector_id=sector_id)
    sectores = Sector.objects.all().order_by('nome')
    return render(request, 'despesas/despesas_fixas_list.html', {'despesas': qs.order_by('sector','categoria'), 'sectores': sectores, 'sector_id': sector_id})

@login_required
@acesso_restrito(['admin'])
def adicionar_despesa_fixa(request):
    if request.method == 'POST':
        form = DespesaFixaForm(request.POST)
        if form.is_valid():
            df = form.save(commit=False)
            if not df.responsavel:
                df.responsavel = request.user
            df.save()
            messages.success(request, '✅ Despesa fixa adicionada.')
            return redirect('listar_despesas_fixas')
        else:
            messages.error(request, '❌ Formulário inválido.')
    else:
        form = DespesaFixaForm()
    return render(request, 'despesas/despesa_fixa_form.html', {'form': form, 'acao': 'Adicionar'})

@login_required
@acesso_restrito(['admin'])
def editar_despesa_fixa(request, pk):
    df = get_object_or_404(DespesaFixa, pk=pk)
    if request.method == 'POST':
        form = DespesaFixaForm(request.POST, instance=df)
        if form.is_valid():
            df = form.save(commit=False)
            df.responsavel = request.user
            df.save()
            messages.success(request, '✅ Despesa fixa atualizada.')
            return redirect('listar_despesas_fixas')
        else:
            messages.error(request, '❌ Formulário inválido.')
    else:
        form = DespesaFixaForm(instance=df)
    return render(request, 'despesas/despesa_fixa_form.html', {'form': form, 'acao': 'Editar', 'despesa': df})

@login_required
@acesso_restrito(['admin'])
def deletar_despesa_fixa(request, pk):
    df = get_object_or_404(DespesaFixa, pk=pk)
    if request.method == 'POST':
        df.delete()
        messages.success(request, '✅ Despesa fixa eliminada.')
        return redirect('listar_despesas_fixas')
    return render(request, 'des/despesa_fixa_confirm_delete.html', {'despesa': df})


# ---------- Depósitos Views ----------#

@login_required
@acesso_restrito(['admin'])
def depositos_view(request):
    """
    Página com abas: Registrar depósito, Listar depósitos, Totais por sector.
    """
    sectores = Sector.objects.all().order_by('nome')
    ultimos = Deposito.objects.select_related('sector','responsavel').order_by('-data_deposito')[:20]
    return render(request, 'depositos/depositos.html', {'sectores': sectores, 'ultimos': ultimos})


@login_required
@require_POST
@acesso_restrito(['admin'])
def depositos_save(request):
    """
    Salvar depósito via POST JSON ou form-POST.
    Aceita JSON no body ou form-data.
    Se for JSON retorna JSON; se for form-POST redireciona para a listagem.
    """
    # detectar JSON independentemente de charset
    content_type = (request.META.get('CONTENT_TYPE') or request.content_type or '').lower()
    is_json = content_type.startswith('application/json')

    if is_json:
        try:
            data = json.loads(request.body.decode('utf-8'))
        except Exception:
            return HttpResponseBadRequest('JSON inválido')
        sector_id = data.get('sector_id')
        data_deposito = data.get('data_deposito')
        valor = data.get('valor')
        observacao = data.get('observacao', '')
    else:
        sector_id = request.POST.get('sector_id')
        data_deposito = request.POST.get('data_deposito')
        valor = request.POST.get('valor')
        observacao = request.POST.get('observacao', '')

    if not sector_id or not valor:
        if is_json:
            return JsonResponse({'ok': False, 'error': 'sector_id e valor obrigatórios'}, status=400)
        messages.error(request, 'sector_id e valor obrigatórios')
        return redirect('depositos_view')

    try:
        sector = Sector.objects.get(pk=sector_id)
    except Sector.DoesNotExist:
        if is_json:
            return JsonResponse({'ok': False, 'error': 'Sector não encontrado'}, status=404)
        messages.error(request, 'Sector não encontrado')
        return redirect('depositos_view')

    try:
        valor_dec = Decimal(str(valor))
    except Exception:
        if is_json:
            return JsonResponse({'ok': False, 'error': 'Valor inválido'}, status=400)
        messages.error(request, 'Valor inválido')
        return redirect('depositos_view')

    dep = Deposito.objects.create(
        sector=sector,
        data_deposito=data_deposito or None,
        valor=valor_dec,
        observacao=observacao,
        responsavel=request.user
    )

    if is_json:
        return JsonResponse({
            'ok': True,
            'deposito_id': dep.id,
            'valor': str(dep.valor),
            'data_deposito': dep.data_deposito.isoformat() if dep.data_deposito else None
        })
    else:
        messages.success(request, '✅ Depósito adicionado com sucesso!')
        return redirect('depositos_view')


@login_required
@acesso_restrito(['admin'])
def depositos_list(request):
    """
    API para listar depósitos com filtros opcionais:
    GET params: sector_id, year, month, limit
    """
    qs = Deposito.objects.select_related('sector','responsavel').all()

    # ✅ FILTRO POR SECTOR
    sector_id = request.GET.get('sector_id')
    if sector_id:
        qs = qs.filter(sector_id=sector_id)

    # ✅ FILTRO POR MÊS E ANO
    year = request.GET.get('year')
    month = request.GET.get('month')
    if year and month:
        qs = qs.filter(data_deposito__year=year,
                       data_deposito__month=month)

    # Limite opcional
    limit = request.GET.get('limit')
    if limit:
        try:
            qs = qs[:int(limit)]
        except:
            pass

    # Total filtrado
    total = qs.aggregate(total_valor=Sum('valor'))['total_valor'] or Decimal('0.00')

    items = []
    for d in qs.order_by('-data_deposito')[:200]:
        items.append({
            'id': d.id,
            'sector_id': d.sector_id,
            'sector': d.sector.nome,
            'data_deposito': d.data_deposito.isoformat() if d.data_deposito else None,
            'valor': str(d.valor),
            'observacao': d.observacao or '',
            'responsavel': d.responsavel.get_full_name() if d.responsavel else ''
        })

    return JsonResponse({
        'ok': True,
        'total': str(total),
        'depositos': items
    })


@login_required
@acesso_restrito(['admin'])
def depositos_detail(request, pk):
    """Retorna um depósito específico em JSON"""
    try:
        d = Deposito.objects.select_related('sector','responsavel').get(pk=pk)
    except Deposito.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Depósito não encontrado'}, status=404)
    return JsonResponse({'ok': True, 'deposito': {
        'id': d.id,
        'sector_id': d.sector_id,
        'sector': d.sector.nome,
        'data_deposito': d.data_deposito.isoformat() if d.data_deposito else None,
        'valor': str(d.valor),
        'observacao': d.observacao or '',
        'responsavel': d.responsavel.get_full_name() if d.responsavel else ''
    }})


@login_required
@require_POST
@acesso_restrito(['admin'])
def depositos_edit(request, pk):
    """
    Edita depósito via POST JSON ou form.
    Campos aceitos: data_deposito, valor, observacao, sector_id
    """
    try:
        dep = Deposito.objects.get(pk=pk)
    except Deposito.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Depósito não encontrado'}, status=404)

    content_type = (request.META.get('CONTENT_TYPE') or request.content_type or '').lower()
    if content_type.startswith('application/json'):
        try:
            data = json.loads(request.body.decode('utf-8'))
        except Exception:
            return HttpResponseBadRequest('JSON inválido')
        data_deposito = data.get('data_deposito')
        valor = data.get('valor')
        observacao = data.get('observacao', dep.observacao)
        sector_id = data.get('sector_id')
    else:
        data_deposito = request.POST.get('data_deposito')
        valor = request.POST.get('valor')
        observacao = request.POST.get('observacao', dep.observacao)
        sector_id = request.POST.get('sector_id')

    if sector_id:
        try:
            dep.sector = Sector.objects.get(pk=sector_id)
        except Sector.DoesNotExist:
            return JsonResponse({'ok': False, 'error': 'Sector não encontrado'}, status=404)

    if valor:
        try:
            dep.valor = Decimal(str(valor))
        except Exception:
            return JsonResponse({'ok': False, 'error': 'Valor inválido'}, status=400)
    if data_deposito:
        dep.data_deposito = data_deposito
    dep.observacao = observacao
    dep.responsavel = request.user
    dep.save()
    return JsonResponse({'ok': True, 'deposito_id': dep.id})


@login_required
@require_POST
@acesso_restrito(['admin'])
def depositos_delete(request):
    """
    Excluir depósito via POST JSON {id: pk} ou form (id).
    """
    content_type = (request.META.get('CONTENT_TYPE') or request.content_type or '').lower()
    if content_type.startswith('application/json'):
        try:
            data = json.loads(request.body.decode('utf-8'))
        except Exception:
            return HttpResponseBadRequest('JSON inválido')
        pk = data.get('id')
    else:
        pk = request.POST.get('id')

    if not pk:
        return JsonResponse({'ok': False, 'error': 'id obrigatório'}, status=400)
    try:
        d = Deposito.objects.get(pk=pk)
        d.delete()
    except Deposito.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Depósito não encontrado'}, status=404)
    return JsonResponse({'ok': True})



# 🔹 Dashboards Especializados
from django.db.models import ExpressionWrapper
@login_required
@acesso_restrito(['admin'])
def decimal_default(obj):
    if isinstance(obj, Decimal):
        return float(obj)
    raise TypeError


@login_required
@acesso_restrito(['admin'])
def contabilista_financas(request):
    registos = RegistoDiario.objects.annotate(
        saldo_liquido=ExpressionWrapper(
            (F("normal") + F("alunos") + F("luvu") + F("frete")) -
            (F("alimentacao") + F("parqueamento") + F("taxa") + F("outros")),
            output_field=DecimalField(max_digits=12, decimal_places=2)
        )
    )

    totais = registos.aggregate(
        total_entradas=Sum(F("normal") + F("alunos") + F("luvu") + F("frete"),
                           output_field=DecimalField()),
        total_saidas=Sum(F("alimentacao") + F("parqueamento") + F("taxa") + F("outros"),
                         output_field=DecimalField()),
        total_saldo=Sum("saldo_liquido"),
    )

    # Agregar despesas gerais (Despesas) no sistema
    total_despesas_gerais = Despesa.objects.aggregate(total=Sum('valor', output_field=DecimalField()))['total'] or Decimal('0')

    # Agregar despesas de combustível no sistema (inclui sobragem e lavagem)
    comb_glob = DespesaCombustivel.objects.aggregate(
        total_valor=Sum('valor', output_field=DecimalField()),
        total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
        total_lavagem=Sum('lavagem', output_field=DecimalField()),
    )
    total_combustivel_glob = comb_glob.get('total_valor') or Decimal('0')
    total_combustivel_sobragem_glob = comb_glob.get('total_sobragem') or Decimal('0')
    total_combustivel_lavagem_glob = comb_glob.get('total_lavagem') or Decimal('0')

    # Ajustar total de saídas para incluir despesas gerais e combustível (valor + sobragem + lavagem)
    try:
        orig_saidas = Decimal(totais.get('total_saidas') or 0)
    except Exception:
        orig_saidas = Decimal('0')
    totais['total_saidas'] = orig_saidas + total_despesas_gerais + total_combustivel_glob + total_combustivel_sobragem_glob + total_combustivel_lavagem_glob

    # Recalcular saldo global (entradas - saídas ajustadas)
    try:
        totais['total_entradas'] = Decimal(totais.get('total_entradas') or 0)
    except Exception:
        totais['total_entradas'] = Decimal('0')
    totais['total_saldo'] = totais['total_entradas'] - totais['total_saidas']

    despesas = Despesa.objects.all().order_by("-data")[:10]

    return render(request, "dashboards/contabilista_financas.html", {
        "totais": totais,
        "despesas": despesas,
    })


# ...existing code...
@login_required
@acesso_restrito(['admin'])
def gerencia_financas(request):

    # Captura parâmetro ?mes=YYYY-MM (opcional)
    mes_param = request.GET.get('mes', '').strip()
    ano_mes = None
    if mes_param:
        try:
            ano, mes = map(int, mes_param.split('-'))
            ano_mes = (ano, mes)
        except Exception:
            ano_mes = None

    # =========================== #
    #   AGRUPAMENTOS MENSAIS      #
    # =========================== #
    if ano_mes:
        ano, mes = ano_mes
        registros = (
            RegistoDiario.objects
            .filter(data__year=ano, data__month=mes)
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(
                total_normal=Sum('normal', output_field=DecimalField()),
                total_alunos=Sum('alunos', output_field=DecimalField()),
                total_luvu=Sum('luvu', output_field=DecimalField()),
                total_frete=Sum('frete', output_field=DecimalField()),
                total_alimentacao=Sum('alimentacao', output_field=DecimalField()),
                total_parqueamento=Sum('parqueamento', output_field=DecimalField()),
                total_taxa=Sum('taxa', output_field=DecimalField()),
                total_outros=Sum('outros', output_field=DecimalField()),
            )
            .order_by('mes')
        )

        despesas_relatorio = (
            RelatorioSector.objects
            .filter(data__year=ano, data__month=mes)
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(total_despesas_geral=Sum('despesa_geral', output_field=DecimalField()))
            .order_by('mes')
        )

        combustiveis_qs = (
            DespesaCombustivel.objects
            .filter(data__year=ano, data__month=mes)
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(
                total_valor=Sum('valor', output_field=DecimalField()),
                total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
                total_lavagem=Sum('lavagem', output_field=DecimalField())
            )
            .order_by('mes')
        )

        despesas_fixas = (
            DespesaFixa.objects
            .filter(ativo=True, data_inicio__year=ano, data_inicio__month=mes)
            .annotate(mes=TruncMonth('data_inicio'))
            .values('mes', 'categoria')
            .annotate(total_valor=Sum('valor', output_field=DecimalField()))
            .order_by('mes', 'categoria')
        )

        despesas_variaveis = (
            Despesa.objects
            .filter(data__year=ano, data__month=mes)
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(total_despesas_variaveis=Sum('valor', output_field=DecimalField()))
            .order_by('mes')
        )

    else:
        registros = (
            RegistoDiario.objects
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(
                total_normal=Sum('normal', output_field=DecimalField()),
                total_alunos=Sum('alunos', output_field=DecimalField()),
                total_luvu=Sum('luvu', output_field=DecimalField()),
                total_frete=Sum('frete', output_field=DecimalField()),
                total_alimentacao=Sum('alimentacao', output_field=DecimalField()),
                total_parqueamento=Sum('parqueamento', output_field=DecimalField()),
                total_taxa=Sum('taxa', output_field=DecimalField()),
                total_outros=Sum('outros', output_field=DecimalField()),
            )
            .order_by('mes')
        )

        despesas_relatorio = (
            RelatorioSector.objects
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(total_despesas_geral=Sum('despesa_geral', output_field=DecimalField()))
            .order_by('mes')
        )

        combustiveis_qs = (
            DespesaCombustivel.objects
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(
                total_valor=Sum('valor', output_field=DecimalField()),
                total_sobragem=Sum('sobragem_filtros', output_field=DecimalField()),
                total_lavagem=Sum('lavagem', output_field=DecimalField())
            )
            .order_by('mes')
        )

        despesas_fixas = (
            DespesaFixa.objects
            .filter(ativo=True)
            .annotate(mes=TruncMonth('data_inicio'))
            .values('mes', 'categoria')
            .annotate(total_valor=Sum('valor', output_field=DecimalField()))
            .order_by('mes', 'categoria')
        )

        despesas_variaveis = (
            Despesa.objects
            .annotate(mes=TruncMonth('data'))
            .values('mes')
            .annotate(total_despesas_variaveis=Sum('valor', output_field=DecimalField()))
            .order_by('mes')
        )

    # ============================================
    #   MAPAS PARA ACESSO RÁPIDO POR MÊS
    # ============================================
    despesas_rel_map = {d['mes']: (d.get('total_despesas_geral') or 0) for d in despesas_relatorio}
    despesas_var = {d['mes']: (d.get('total_despesas_variaveis') or 0) for d in despesas_variaveis}

    despesas_fixas_map = {
        (d['mes'], d['categoria']): (d.get('total_valor') or 0)
        for d in despesas_fixas
    }

    categorias_fixas = [c[0] for c in DespesaFixa.CATEGORIAS]

    combustivel_map_valor = {c['mes']: c.get('total_valor') or 0 for c in combustiveis_qs}
    combustivel_map_sobragem = {c['mes']: c.get('total_sobragem') or 0 for c in combustiveis_qs}
    combustivel_map_lavagem = {c['mes']: c.get('total_lavagem') or 0 for c in combustiveis_qs}

    # ================================
    #   LABELS DOS GRÁFICOS
    # ================================
    labels = [r['mes'].strftime("%B %Y") for r in registros]

    # ================================
    #   ENTRADAS
    # ================================
    serie_normal = [float(r['total_normal'] or 0) for r in registros]
    serie_alunos = [float(r['total_alunos'] or 0) for r in registros]
    serie_luvu = [float(r['total_luvu'] or 0) for r in registros]
    serie_frete = [float(r['total_frete'] or 0) for r in registros]

    # Soma correta das entradas por mês
    serie_entradas = [
        serie_normal[i] + serie_alunos[i] + serie_luvu[i] + serie_frete[i]
        for i in range(len(registros))
    ]

    # ================================
    #   SAÍDAS
    # ================================
    serie_alimentacao = [float(r['total_alimentacao'] or 0) for r in registros]
    serie_parqueamento = [float(r['total_parqueamento'] or 0) for r in registros]
    serie_taxa = [float(r['total_taxa'] or 0) for r in registros]
    serie_outros = [float(r['total_outros'] or 0) for r in registros]
    serie_taxi = [float(r['total_taxi'] or 0) for r in registros]
    
    serie_alimentacao_estaleiro = [float(r['total_alimentacao_estaleiro'] or 0) for r in registros]
    serie_despesas_extra = [float(despesas_rel_map.get(r['mes'], 0)) for r in registros]

    serie_combustivel_valor = [float(combustivel_map_valor.get(r['mes'], 0)) for r in registros]
    serie_combustivel_sobragem = [float(combustivel_map_sobragem.get(r['mes'], 0)) for r in registros]
    serie_combustivel_lavagem = [float(combustivel_map_lavagem.get(r['mes'], 0)) for r in registros]

    serie_saidas = [
        (
            serie_alimentacao[i] +
            serie_parqueamento[i] +
            serie_taxa[i] +
            serie_outros[i] +
            serie_taxi[i] +
            serie_despesas_extra[i] +
            serie_alimentacao_estaleiro[i] +
            serie_combustivel_valor[i] +
            serie_combustivel_sobragem[i] +
            serie_combustivel_lavagem[i]
        )
        for i in range(len(registros))
    ]

    # ================================
    #   SALDO MENSAL
    # ================================
    serie_saldo = [
        serie_entradas[i] - serie_saidas[i]
        for i in range(len(registros))
    ]

    # ================================
    #   DESPESAS FIXAS (por categoria)
    # ================================
    serie_despesas_fixas = {
        categoria: [
            float(despesas_fixas_map.get((r['mes'], categoria), 0))
            for r in registros
        ]
        for categoria in categorias_fixas
    }

    # Soma total das despesas fixas por mês
    serie_total_despesas_fixas = [
        sum(serie_despesas_fixas[c][i] for c in categorias_fixas)
        for i in range(len(registros))
    ]

    # ================================
    #   DESPESAS VARIÁVEIS
    # ================================
    serie_despesas_variaveis = [
        float(despesas_var.get(r['mes'], 0)) for r in registros
    ]

    # ================================
    # TOTAL DE DESPESAS
    # ================================
    # (fixas + variáveis)
    serie_total_despesas = [
        serie_total_despesas_fixas[i] + serie_despesas_variaveis[i]
        for i in range(len(registros))
    ]

    # ================================
    #   LUCRO MENSAL
    # ================================
    serie_lucro = [
        serie_saldo[i] - serie_total_despesas_fixas[i] - serie_despesas_variaveis[i]
        for i in range(len(registros))
    ]

    # ================================
    #   CONTEXTO FINAL
    # ================================
    context = {
        "labels": labels,

        # Entradas
        "serie_normal": serie_normal,
        "serie_alunos": serie_alunos,
        "serie_luvu": serie_luvu,
        "serie_frete": serie_frete,
        "serie_entradas": serie_entradas,

        # Saídas
        "serie_alimentacao": serie_alimentacao,
        "serie_parqueamento": serie_parqueamento,
        "serie_taxa": serie_taxa,
        "serie_outros": serie_outros,
        "serie_taxi": serie_taxi,
        "serie_despesas_extra": serie_despesas_extra,
        "serie_alimentacao_estaleiro": serie_alimentacao_estaleiro,

        # Combustível
        "serie_combustivel_valor": serie_combustivel_valor,
        "serie_combustivel_sobragem": serie_combustivel_sobragem,
        "serie_combustivel_lavagem": serie_combustivel_lavagem,

        # Saidas Diarias E Semanais
        "serie_saidas": serie_saidas,

        # Saldo
        "serie_saldo": serie_saldo,

        # Despesas fixas
        "serie_total_despesas_fixas": serie_total_despesas_fixas,

        # Despesas variáveis
        "serie_despesas_variaveis": serie_despesas_variaveis,

        # Total De Despesa
        "serie_total_despesas": serie_total_despesas,

        # Lucro final
        "serie_lucro": serie_lucro,
    }

    # ---------------------------
    # download CSV com TODOS os dados exibidos no template
    # ---------------------------
    """if request.GET.get('download') == '1':
        import csv
        from django.http import HttpResponse

        def _safe(series, i):
            try:
                return float(series[i]) if series[i] is not None else 0.0
            except Exception:
                return 0.0

        filename = f"relatorio_financas_{mes_param or 'todos_meses'}.csv"
        resp = HttpResponse(content_type='text/csv')
        resp['Content-Disposition'] = f'attachment; filename="{filename}"'
        writer = csv.writer(resp)

        header = [
            "Mês",
            "Normal", "Alunos", "Luvu", "Frete", "Entradas",
            "Alimentação", "Taxas", "Outros", "Parqueamento", "Despesas Extra (Relatorio)",
            "Combustível (valor)", "Sobragem/Filtros", "Lavagem",
            "Saídas (total)",
            "Saldo",
            "Despesas Fixas (total)", "Despesas Variáveis",
            "Lucro"
        ]
        writer.writerow(header)

        for i, label in enumerate(labels):
            row = [
                label,
                f"{_safe(serie_normal, i):.2f}",
                f"{_safe(serie_alunos, i):.2f}",
                f"{_safe(serie_luvu, i):.2f}",
                f"{_safe(serie_frete, i):.2f}",
                f"{_safe(serie_entradas, i):.2f}",
                f"{_safe(serie_alimentacao, i):.2f}",
                f"{_safe(serie_taxa, i):.2f}",
                f"{_safe(serie_outros, i):.2f}",
                f"{_safe(serie_parqueamento, i):.2f}",
                f"{_safe(serie_despesas_extra, i):.2f}",
                f"{_safe(serie_combustivel_valor, i):.2f}",
                f"{_safe(serie_combustivel_sobragem, i):.2f}",
                f"{_safe(serie_combustivel_lavagem, i):.2f}",
                f"{_safe(serie_saidas, i):.2f}",
                f"{_safe(serie_saldo, i):.2f}",
                f"{_safe(serie_total_despesas_fixas, i):.2f}",
                f"{_safe(serie_despesas_variaveis, i):.2f}",
                f"{_safe(serie_lucro, i):.2f}",
            ]
            writer.writerow(row)

        return resp"""

    # ---------------------------
    # download DOCX com os dados do mês selecionado (ou último mês disponível)
    # ---------------------------
    if request.GET.get('download') == '1':
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        def _fmt_kz(value):
            try:
                v = Decimal(value or 0).quantize(Decimal('0.01'))
                s = f"{v:,.2f}"               # ex: 1,234.56
                s = s.replace(',', 'X').replace('.', ',').replace('X', '.')
                return f"{s} Kz"
            except Exception:
                return "0,00 Kz"

        # garantir que pegamos um único mês: se view já filtrou por mes_param, registros terá apenas esse mês;
        # caso contrário, usamos o último mês da lista.
        if labels:
            label = labels[-1]
            idx = len(labels) - 1
        else:
            label = mes_param or timezone.now().strftime("%B %Y")
            idx = 0

        doc = Document()
        # paisagem A4-like
        sec = doc.sections[0]
        sec.page_width = Inches(11.69)
        sec.page_height = Inches(8.27)
        sec.left_margin = sec.right_margin = Inches(0.5)
        sec.top_margin = sec.bottom_margin = Inches(0.5)

        # Cabeçalho simples
        hdr = doc.add_paragraph()
        run = hdr.add_run("Relatório Financeiro — " + (mes_param or label))
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph()

        # 1) Entradas (linha)
        tbl_e = doc.add_table(rows=2, cols=5)
        tbl_e.style = 'Table Grid'
        hdr_cells = tbl_e.rows[0].cells
        hdr_cells[0].text = "Normal"
        hdr_cells[1].text = "Alunos"
        hdr_cells[2].text = "Luvu"
        hdr_cells[3].text = "Frete"
        hdr_cells[4].text = "Entradas (Total)"
        row_cells = tbl_e.rows[1].cells
        row_cells[0].text = _fmt_kz(serie_normal[idx] if serie_normal else 0)
        row_cells[1].text = _fmt_kz(serie_alunos[idx] if serie_alunos else 0)
        row_cells[2].text = _fmt_kz(serie_luvu[idx] if serie_luvu else 0)
        row_cells[3].text = _fmt_kz(serie_frete[idx] if serie_frete else 0)
        row_cells[4].text = _fmt_kz(serie_entradas[idx] if serie_entradas else 0)
        for c in range(5):
            try:
                tbl_e.rows[1].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except Exception:
                pass
        doc.add_paragraph()

        # 2) Saídas - primeira linha: alimentacao, taxas, outros, parqueamento
        tbl_s1 = doc.add_table(rows=2, cols=4)
        tbl_s1.style = 'Table Grid'
        hdr = tbl_s1.rows[0].cells
        hdr[0].text = "Alimentação"
        hdr[1].text = "Taxas"
        hdr[2].text = "Outros"
        hdr[3].text = "Parqueamento"
        vals = tbl_s1.rows[1].cells
        vals[0].text = _fmt_kz(serie_alimentacao[idx] if serie_alimentacao else 0)
        vals[1].text = _fmt_kz(serie_taxa[idx] if serie_taxa else 0)
        vals[2].text = _fmt_kz(serie_outros[idx] if serie_outros else 0)
        vals[3].text = _fmt_kz(serie_parqueamento[idx] if serie_parqueamento else 0)
        for c in range(4):
            try:
                tbl_s1.rows[1].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except Exception:
                pass
        doc.add_paragraph()

        # 3) Saídas - segunda linha: Despesas Extra, Combustível, Lavagem, Sobragem
        tbl_s2 = doc.add_table(rows=2, cols=4)
        tbl_s2.style = 'Table Grid'
        hdr = tbl_s2.rows[0].cells
        hdr[0].text = "Despesas Extra"
        hdr[1].text = "Combustível (valor)"
        hdr[2].text = "Lavagem"
        hdr[3].text = "Sobragem / Filtros"
        vals = tbl_s2.rows[1].cells
        vals[0].text = _fmt_kz(serie_despesas_extra[idx] if serie_despesas_extra else 0)
        vals[1].text = _fmt_kz(serie_combustivel_valor[idx] if serie_combustivel_valor else 0)
        vals[2].text = _fmt_kz(serie_combustivel_lavagem[idx] if serie_combustivel_lavagem else 0)
        vals[3].text = _fmt_kz(serie_combustivel_sobragem[idx] if serie_combustivel_sobragem else 0)
        for c in range(4):
            try:
                tbl_s2.rows[1].cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except Exception:
                pass
        doc.add_paragraph()

        # 4) Saldo (linha única)
        tbl_saldo = doc.add_table(rows=1, cols=2)
        tbl_saldo.style = 'Table Grid'
        tbl_saldo.rows[0].cells[0].text = "Saldo (Entradas - Saídas)"
        tbl_saldo.rows[0].cells[1].text = _fmt_kz(serie_saldo[idx] if serie_saldo else 0)
        try:
            tbl_saldo.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass
        doc.add_paragraph()

        # 5) Despesas: Despesas Fixas e Variáveis (uma linha com duas colunas)
        tbl_d = doc.add_table(rows=1, cols=2)
        tbl_d.style = 'Table Grid'
        tbl_d.rows[0].cells[0].text = "Despesas Fixas (Total)"
        tbl_d.rows[0].cells[1].text = _fmt_kz(serie_total_despesas_fixas[idx] if serie_total_despesas_fixas else 0)
        try:
            tbl_d.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass
        doc.add_paragraph()
        tbl_dv = doc.add_table(rows=1, cols=2)
        tbl_dv.style = 'Table Grid'
        tbl_dv.rows[0].cells[0].text = "Despesas Variáveis (Total)"
        tbl_dv.rows[0].cells[1].text = _fmt_kz(serie_despesas_variaveis[idx] if serie_despesas_variaveis else 0)
        try:
            tbl_dv.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass
        doc.add_paragraph()

        # 6) Lucro (linha única)
        tbl_l = doc.add_table(rows=1, cols=2)
        tbl_l.style = 'Table Grid'
        tbl_l.rows[0].cells[0].text = "Lucro Final"
        tbl_l.rows[0].cells[1].text = _fmt_kz(serie_lucro[idx] if serie_lucro else 0)
        try:
            tbl_l.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass

        # enviar DOCX
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"Relatorio_Financas_{mes_param or label}.docx"
        from django.http import HttpResponse
        resp = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = f'attachment; filename="{filename}"'
        return resp

    return render(request, "dashboards/gerencia_financas.html", context)


@login_required
@acesso_restrito(['admin'])
def gerencia_campo(request):
    # Verificar se o modelo Motorista existe
    try:
        from .models import Motorista
        MOTORISTA_MODEL_EXISTS = True
    except ImportError:
        MOTORISTA_MODEL_EXISTS = False

    if hasattr(Autocarro, "status"):
        autocarros_ativos = Autocarro.objects.filter(status="ativo").count()
        autocarros_inativos = Autocarro.objects.filter(status="inativo").count()
        autocarros_manutencao = Autocarro.objects.filter(status="manutencao").count()
        autocarros_queryset = Autocarro.objects.all().order_by("numero")
    else:
        autocarros_ativos = Autocarro.objects.count()
        autocarros_inativos = Autocarro.objects.count()
        autocarros_manutencao = 0
        autocarros_queryset = Autocarro.objects.all().order_by("numero")

    if MOTORISTA_MODEL_EXISTS:
        motoristas_ativos = Motorista.objects.filter(ativo=True).count()
    else:
        motoristas_ativos = RegistoDiario.objects.exclude(motorista__isnull=True).exclude(motorista__exact="").values("motorista").distinct().count()

    autocarros_map = autocarros_queryset.values("pk", "numero", "modelo", "lat", "lng", "status")

    context = {
        "autocarros_ativos": autocarros_ativos,
        "autocarros_inativos": autocarros_inativos,
        "autocarros_manutencao": autocarros_manutencao,
        "motoristas_ativos": motoristas_ativos,
        "autocarros": list(autocarros_map),
    }
    return render(request, "dashboards/gerencia_campo.html", context)


# ---------- Cobrador Viagens Views ----------#

@login_required
@acesso_restrito(['admin'])
def cobrador_viagens(request):
    """
    Página com formulário/abas para lançar e ver viagens do cobrador.
    Template: templates/cobradores/cobrador_viagens.html
    """
    # enviar data inicial se quiser
    context = {
        "today": request.GET.get("data") or None
    }
    return render(request, "cobradores/cobrador_viagens.html", context)

@login_required
@require_POST
def cobrador_viagens_save(request):
    """
    Salva uma viagem via POST JSON.
    Espera JSON no body com: autocarro_numero | autocarro_id, data, hora, valor, passageiros, observacao
    """
    try:
        data = json.loads(request.body.decode("utf-8"))
    except Exception:
        return HttpResponseBadRequest("JSON inválido")

    numero = data.get("autocarro_numero") or data.get("autocarro_id")
    if not numero:
        return JsonResponse({"ok": False, "error": "autocarro_numero obrigatório"}, status=400)

    # localizar autocarro (tenta por número, depois por pk)
    try:
        try:
            autocarro = Autocarro.objects.get(numero=numero)
        except Autocarro.DoesNotExist:
            autocarro = Autocarro.objects.get(pk=int(numero))
    except Exception:
        return JsonResponse({"ok": False, "error": "Autocarro não encontrado"}, status=404)

    try:
        valor = Decimal(str(data.get("valor") or "0"))
    except Exception:
        valor = Decimal("0")

    try:
        passageiros = int(data.get("passageiros") or 0)
    except Exception:
        passageiros = 0

    viagem = CobradorViagem.objects.create(
        autocarro=autocarro,
        cobrador=request.user,
        data=data.get("data"),
        hora=data.get("hora") or None,
        valor=valor,
        passageiros=passageiros,
        observacao=data.get("observacao") or ""
    )

    return JsonResponse({"ok": True, "viagem_id": viagem.id})


@login_required
def cobrador_viagens_list(request):
    """
    Retorna JSON com viagens e resumo para um autocarro (por número ou id).
    GET params: autocarro_numero|autocarro_id, opcional data=YYYY-MM-DD
    """
    numero = request.GET.get("autocarro_numero") or request.GET.get("autocarro_id")
    if not numero:
        return JsonResponse({"ok": False, "error": "autocarro_numero obrigatório"}, status=400)

    try:
        try:
            autocarro = Autocarro.objects.get(numero=numero)
        except Autocarro.DoesNotExist:
            autocarro = Autocarro.objects.get(pk=int(numero))
    except Exception:
        return JsonResponse({"ok": False, "error": "Autocarro não encontrado"}, status=404)

    qs = CobradorViagem.objects.filter(autocarro=autocarro)
    data_filtro = request.GET.get("data")
    if data_filtro:
        qs = qs.filter(data=data_filtro)

    viagens = []
    for v in qs.order_by("-data", "-hora"):
        viagens.append({
            "id": v.id,
            "data": v.data.isoformat(),
            "hora": v.hora.isoformat() if v.hora else "",
            "valor": str(v.valor),
            "passageiros": v.passageiros,
            "observacao": v.observacao,
            "cobrador": v.cobrador.get_full_name() if v.cobrador else "",
            "status": getattr(v, "status", "pending"),
        })

    resumo = qs.aggregate(total_valor=Sum("valor"), total_passageiros=Sum("passageiros"))
    total_valor = resumo.get("total_valor") or Decimal("0")
    total_passageiros = resumo.get("total_passageiros") or 0

    return JsonResponse({
        "ok": True,
        "autocarro": {"id": autocarro.id, "numero": autocarro.numero, "modelo": getattr(autocarro, "modelo", "")},
        "viagens": viagens,
        "resumo": {"total_valor": str(total_valor), "total_passageiros": int(total_passageiros)}
    })


@login_required
def cobrador_viagens_validate_list(request):
    """
    Retorna viagens pendentes (ou filtro por autocarro_numero/data).
    Apenas para users com nível 'gestor' ou 'admin'.
    """
    nivel = getattr(request.user, 'nivel_acesso', '').lower()
    if nivel not in ['admin', 'gestor']:
        return JsonResponse({'ok': False, 'error': 'Acesso negado'}, status=403)

    numero = request.GET.get('autocarro_numero') or request.GET.get('autocarro_id')
    data = request.GET.get('data')  # opcional
    qs = CobradorViagem.objects.filter(status='pending')
    if numero:
        try:
            aut = Autocarro.objects.get(numero=numero)
            qs = qs.filter(autocarro=aut)
        except Autocarro.DoesNotExist:
            try:
                qs = qs.filter(autocarro_id=int(numero))
            except Exception:
                return JsonResponse({'ok': False, 'error': 'Autocarro não encontrado'}, status=404)
    if data:
        qs = qs.filter(data=data)
    viagens = []
    for v in qs.order_by('data', 'hora'):
        viagens.append({
            'id': v.id,
            'autocarro_numero': v.autocarro.numero,
            'data': v.data.isoformat(),
            'hora': v.hora.isoformat() if v.hora else '',
            'valor': str(v.valor),
            'passageiros': v.passageiros,
            'observacao': v.observacao,
            'cobrador': v.cobrador.get_full_name() if v.cobrador else '',
            'criado_em': v.criado_em.isoformat(),
        })
    return JsonResponse({'ok': True, 'viagens': viagens})


@login_required
@require_POST
def cobrador_viagens_validate_action(request):
    """
    Ação para aprovar/reprovar via POST JSON:
    { "id": 123, "action": "approve"|"reject", "valor_aprovado": "1000.00", "nota": "..." }
    Apenas 'admin' e 'gestor' podem executar.
    """
    nivel = getattr(request.user, 'nivel_acesso', '').lower()
    if nivel not in ['admin', 'gestor']:
        return JsonResponse({'ok': False, 'error': 'Acesso negado'}, status=403)
    try:
        data = json.loads(request.body.decode('utf-8'))
    except Exception:
        return HttpResponseBadRequest('JSON inválido')
    vid = data.get('id')
    action = data.get('action')
    nota = data.get('nota')
    valor_aprovado = data.get('valor_aprovado', None)
    if not vid or action not in ['approve', 'reject']:
        return JsonResponse({'ok': False, 'error': 'Parâmetros inválidos'}, status=400)
    try:
        viagem = CobradorViagem.objects.get(pk=vid)
    except CobradorViagem.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Viagem não encontrada'}, status=404)

    try:
        if action == 'approve':
            viagem.approve(request.user, valor_aprovado=valor_aprovado, nota=nota)
        else:
            viagem.reject(request.user, nota=nota)
        return JsonResponse({'ok': True, 'status': viagem.status})
    except Exception as e:
        return JsonResponse({'ok': False, 'error': str(e)}, status=500)


# ---------- Manutenção Autocarros Views ----------#
@login_required
def manutencao_create(request):
    if request.method == 'POST':
        form = ManutencaoForm(request.POST)
        if form.is_valid():
            m = form.save(commit=False)

            # Define automaticamente o responsável
            if not m.responsavel:
                m.responsavel = request.user

            # Força o status inicial a "agendada"
            m.status = 'agendada'

            # Campos automáticos (km_proxima e km_prox_*) são calculados no save()
            m.save()

            messages.success(request, '✅ Manutenção agendada com sucesso!')
            return redirect('manutencao_list')
        else:
            for field, errors in form.errors.items():
                for e in errors:
                    messages.error(request, f"{field}: {e}")

    else:
        # inicializa o formulário
        form = ManutencaoForm(initial={'status': 'agendada'})

    sectores = Sector.objects.all().order_by('nome')

    return render(request, 'autocarros/manutencao_form.html', {
        'form': form,
        'sectores': sectores,
    })


@login_required
def manutencao_list(request):
    qs = Manutencao.objects.select_related('autocarro', 'sector', 'responsavel').all()

    # Tentar ordenar por '-data_ultima' e, se não existir no modelo, usar um fallback para '-data' ou sem ordenação.
    order_field = None
    try:
        Manutencao._meta.get_field('data_ultima')
        order_field = '-data_ultima'
    except Exception:
        try:
            Manutencao._meta.get_field('data')
            order_field = '-data'
        except Exception:
            order_field = None

    if order_field:
        qs = qs.order_by(order_field)

    sector_id = request.GET.get('sector')
    if sector_id:
        qs = qs.filter(sector_id=sector_id)

    return render(request, 'autocarros/manutencao_list.html', {'manutencoes': qs, 'sectores': Sector.objects.all()})


@login_required
@acesso_restrito(['admin'])
def manutencao_edit(request, pk):
    manut = get_object_or_404(Manutencao, pk=pk)
    if request.method == 'POST':
        form = ManutencaoForm(request.POST, request.FILES, instance=manut)
        if form.is_valid():
            form.save()
            messages.success(request, '✅ Manutenção atualizada com sucesso.')
            return redirect('manutencao_list')
        else:
            messages.error(request, '❌ Formulário inválido. Verifique os campos.')
    else:
        form = ManutencaoForm(instance=manut)
    return render(request, 'autocarros/manutencao_form.html', {'form': form, 'manut': manut})


@login_required
@acesso_restrito(['admin', 'gestor'])
def manutencao_delete(request, pk):
    manut = get_object_or_404(Manutencao, pk=pk)
    if request.method == 'POST':
        manut.delete()
        messages.success(request, '✅ Manutenção eliminada com sucesso.')
        return redirect('manutencao_list')
    return render(request, 'autocarros/confirmar_deletar_manutencao.html', {'manut': manut})


@login_required
def api_autocarros_por_sector(request):
    sector_id = request.GET.get('sector_id')
    if not sector_id:
        return JsonResponse({'ok': False, 'error': 'sector_id obrigatório'}, status=400)
    autos = Autocarro.objects.filter(sector_id=sector_id).values('id','numero','modelo')
    return JsonResponse({'ok': True, 'autocarros': list(autos)})


@login_required
def registro_km_view(request):
    sectores = Sector.objects.all().order_by('nome')
    # listar últimos registros (paginacao simples: últimos 20)
    registros = RegistroKM.objects.select_related('sector').prefetch_related('itens__autocarro').order_by('-data_registo')[:20]

    # preparar dados para listagem com previsão (km_proxima vs km_atual)
    registros_data = []
    for r in registros:
        itens = []
        for it in r.itens.all():
            # buscar manutenção mais recente para o autocarro
            m = Manutencao.objects.filter(autocarro=it.autocarro).order_by('-data_ultima').first()
            km_prox = m.km_proxima if m else None
            falta = None
            status = 'Sem plano'
            if km_prox is not None:
                falta = int(km_prox) - int(it.km_atual)
                if falta <= 0:
                    status = 'Vencida'
                elif falta <= 1000:
                    status = 'Próxima'
                else:
                    status = 'OK'
            itens.append({
                'autocarro_numero': it.autocarro.numero,
                'km_atual': it.km_atual,
                'km_prox': km_prox,
                'falta': falta,
                'status': status
            })
        registros_data.append({
            'id': r.id,
            'sector': r.sector.nome,
            'data_registo': r.data_registo.isoformat(),
            'itens': itens
        })

    return render(request, 'autocarros/registro_km.html', {
        'sectores': sectores,
        'registros': registros_data
    })


@login_required
@require_POST
def registro_km_save(request):
    try:
        data = json.loads(request.body.decode('utf-8'))
    except Exception:
        return HttpResponseBadRequest('JSON inválido')

    sector_id = data.get('sector_id')
    itens = data.get('itens', [])
    data_registo = data.get('data_registo', None)
    if not sector_id or not isinstance(itens, list) or not itens:
        return JsonResponse({'ok': False, 'error': 'sector_id e itens obrigatórios'}, status=400)

    try:
        sector = Sector.objects.get(pk=sector_id)
    except Sector.DoesNotExist:
        return JsonResponse({'ok': False, 'error': 'Sector não encontrado'}, status=404)

    registro = RegistroKM.objects.create(sector=sector, data_registo=data_registo or None)
    created = 0
    for it in itens:
        try:
            aid = int(it.get('autocarro_id'))
            km = int(it.get('km_atual'))
            aut = Autocarro.objects.get(pk=aid, sector=sector)
            RegistroKMItem.objects.create(registro=registro, autocarro=aut, km_atual=km)
            created += 1
        except Exception:
            continue

    return JsonResponse({'ok': True, 'registro_id': registro.id, 'created': created})




# ---------- Mapa Geral Financeiro View ----------#
from decimal import Decimal
from datetime import timedelta

from django.shortcuts import render
from django.utils.timezone import now
from django.contrib.auth.decorators import login_required

from django.db.models import Sum

from .models import (
    RegistoDiario,
    DespesaCombustivel,
    Sector,
    RelatorioSector,
)
from .decorators import acesso_restrito


# ================================
# FUNÇÃO: SEMANA DO MÊS (4 COLUNAS)
# ================================
"""def semana_do_mes_4colunas(data):
    primeiro_dia = data.replace(day=1)

    # weekday(): segunda=0 ... domingo=6
    if primeiro_dia.weekday() == 6:
        fim_primeira_semana = primeiro_dia
    else:
        fim_primeira_semana = primeiro_dia + timedelta(
            days=(6 - primeiro_dia.weekday())
        )

    if data <= fim_primeira_semana:
        return 1

    dias_apos = (data - fim_primeira_semana).days
    semana = 1 + ((dias_apos - 1) // 7) + 1

    return min(semana, 4)"""

# Resovido problema de primeira semana começar no domingo
def semana_do_mes_4colunas(data):
    primeiro_dia = data.replace(day=1)

    if primeiro_dia.weekday() == 6:  # domingo
        fim_primeira_semana = primeiro_dia + timedelta(days=6)
    else:
        fim_primeira_semana = primeiro_dia + timedelta(
            days=(6 - primeiro_dia.weekday())
        )

    if data <= fim_primeira_semana:
        return 1

    dias_apos = (data - fim_primeira_semana).days
    semana = 2 + ((dias_apos - 1) // 7)

    return min(semana, 4)


# ================================
# MAPA GERAL FINANCEIRO
# ================================
@login_required
@acesso_restrito(['admin'])
def mapa_geral_financeiro(request):
    sector_id = request.GET.get("sector")
    mes = int(request.GET.get("mes", now().month))
    ano = int(request.GET.get("ano", now().year))

    sectores = Sector.objects.all()
    sector = None

    registos = RegistoDiario.objects.filter(
        data__year=ano,
        data__month=mes,
    )

    combustiveis = DespesaCombustivel.objects.filter(
        data__year=ano,
        data__month=mes
    )

    despesas_gerais = RelatorioSector.objects.filter(
        data__year=ano,
        data__month=mes
    )

    if sector_id and sector_id != "all":
        sector = Sector.objects.get(id=sector_id)
        registos = registos.filter(autocarro__sector=sector)
        combustiveis = combustiveis.filter(sector=sector)
        despesas_gerais = despesas_gerais.filter(sector=sector)

    # ================================
    # ESTRUTURA DAS 4 SEMANAS
    # ================================
    semanas = {
        i: {
            "entradas": {
                "normal": Decimal(0),
                "alunos": Decimal(0),
                "luvu": Decimal(0),
                "frete": Decimal(0),
                "total": Decimal(0),
            },
            "despesas": {
                "alimentacao": Decimal(0),
                "parqueamento": Decimal(0),
                "taxa": Decimal(0),
                "outros": Decimal(0),
                "taxi": Decimal(0),
                "combustivel": Decimal(0),
                "lavagem": Decimal(0),
                "sopragem": Decimal(0),
                "despesa_geral": Decimal(0),
                "alimentacao_estaleiro": Decimal(0),
                "total": Decimal(0),
            },
            "saldo": Decimal(0),
        }
        for i in range(1, 5)
    }

    # ================================
    # REGISTOS DIÁRIOS
    # ================================
    for r in registos:
        s = semana_do_mes_4colunas(r.data)

        semanas[s]["entradas"]["normal"] += r.normal
        semanas[s]["entradas"]["alunos"] += r.alunos
        semanas[s]["entradas"]["luvu"] += r.luvu
        semanas[s]["entradas"]["frete"] += r.frete

        semanas[s]["despesas"]["alimentacao"] += r.alimentacao
        semanas[s]["despesas"]["parqueamento"] += r.parqueamento
        semanas[s]["despesas"]["taxa"] += r.taxa
        semanas[s]["despesas"]["outros"] += r.outros
        semanas[s]["despesas"]["taxi"] += r.taxi
    # ================================
    # DESPESAS DE COMBUSTÍVEL
    # ================================
    for d in combustiveis:
        s = semana_do_mes_4colunas(d.data)

        semanas[s]["despesas"]["combustivel"] += d.valor or Decimal(0)
        semanas[s]["despesas"]["lavagem"] += d.lavagem or Decimal(0)
        semanas[s]["despesas"]["sopragem"] += d.sobragem_filtros or Decimal(0)

    # ================================
    # DESPESA GERAL (CORREÇÃO PRINCIPAL)
    # ================================
    for dg in despesas_gerais:
        s = semana_do_mes_4colunas(dg.data)
        semanas[s]["despesas"]["despesa_geral"] += dg.despesa_geral or Decimal(0)
        semanas[s]["despesas"]["alimentacao_estaleiro"] += dg.alimentacao_estaleiro or Decimal(0)

    # ================================
    # TOTAIS E SALDOS
    # ================================
    total_entradas = Decimal(0)
    total_despesas = Decimal(0)

    for s in semanas.values():
        s["entradas"]["total"] = sum(
            v for k, v in s["entradas"].items() if k != "total"
        )
        s["despesas"]["total"] = sum(
            v for k, v in s["despesas"].items() if k != "total"
        )
        s["saldo"] = s["entradas"]["total"] - s["despesas"]["total"]

        total_entradas += s["entradas"]["total"]
        total_despesas += s["despesas"]["total"]

    context = {
        "sectores": sectores,
        "sector": sector,
        "mes": mes,
        "ano": ano,
        "semanas": semanas,

        "total_entradas": total_entradas,
        "total_despesas": total_despesas,
        "saldo_liquido": total_entradas - total_despesas,

        "total_entradas_normal": sum(s["entradas"]["normal"] for s in semanas.values()),
        "total_entradas_alunos": sum(s["entradas"]["alunos"] for s in semanas.values()),
        "total_entradas_luvu": sum(s["entradas"]["luvu"] for s in semanas.values()),
        "total_entradas_frete": sum(s["entradas"]["frete"] for s in semanas.values()),

        "total_despesas_alimentacao": sum(s["despesas"]["alimentacao"] for s in semanas.values()),
        "total_despesas_parqueamento": sum(s["despesas"]["parqueamento"] for s in semanas.values()),
        "total_despesas_taxa": sum(s["despesas"]["taxa"] for s in semanas.values()),
        "total_despesas_combustivel": sum(s["despesas"]["combustivel"] for s in semanas.values()),
        "total_despesas_lavagem": sum(s["despesas"]["lavagem"] for s in semanas.values()),
        "total_despesas_sopragem": sum(s["despesas"]["sopragem"] for s in semanas.values()),
        "total_despesas_outros": sum(s["despesas"]["outros"] for s in semanas.values()),
        "total_despesas_taxi": sum(s["despesas"]["taxi"] for s in semanas.values()),
        "total_despesa_geral": sum(s["despesas"]["despesa_geral"] for s in semanas.values()),
        "total_despesa_alimentacao_estaleiro": sum(s["despesas"]["alimentacao_estaleiro"] for s in semanas.values()),
    }
    return render(request, "financeiro/mapa_geral.html", context)


# ---------- Gestão de Despesas Views ----------#
# =============================
# CADASTRAR CATEGORIA
# =============================
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from django.contrib import messages

from .models import CategoriaDespesa, SubCategoriaDespesa
from .forms import CategoriaDespesaForm, SubCategoriaDespesaForm


# ============================
# CATEGORIA
# ============================
@login_required
@acesso_restrito(['admin'])
def categoria_create(request):
    form = CategoriaDespesaForm(request.POST or None)
    if form.is_valid():
        form.save()
        messages.success(request, 'Categoria cadastrada com sucesso.')
        return redirect('categoria_create')

    categorias = CategoriaDespesa.objects.all()
    return render(request, 'financeiro/categoria_form.html', {
        'form': form,
        'categorias': categorias
    })

@login_required
@acesso_restrito(['admin'])
def categoria_update(request, pk):
    categoria = get_object_or_404(CategoriaDespesa, pk=pk)
    form = CategoriaDespesaForm(request.POST or None, instance=categoria)

    if form.is_valid():
        form.save()
        messages.success(request, 'Categoria atualizada.')
        return redirect('categoria_create')

    return render(request, 'financeiro/categoria_form.html', {
        'form': form,
        'categorias': CategoriaDespesa.objects.all(),
        'editando': True
    })

@login_required
@acesso_restrito(['admin'])
def categoria_delete(request, pk):
    categoria = get_object_or_404(CategoriaDespesa, pk=pk)
    categoria.delete()
    messages.success(request, 'Categoria removida.')
    return redirect('categoria_create')


# ============================
# SUBCATEGORIA
# ============================
@login_required
@acesso_restrito(['admin'])
def subcategoria_create(request):
    categoria_id = request.GET.get('categoria')

    form = SubCategoriaDespesaForm(request.POST or None)
    if form.is_valid():
        form.save()
        messages.success(request, 'Subcategoria cadastrada.')
        return redirect('subcategoria_create')

    subcategorias = SubCategoriaDespesa.objects.select_related('categoria')

    if categoria_id:
        subcategorias = subcategorias.filter(categoria_id=categoria_id)

    categorias = CategoriaDespesa.objects.filter(ativa=True)

    return render(request, 'financeiro/subcategoria_form.html', {
        'form': form,
        'subcategorias': subcategorias,
        'categorias': categorias,
        'categoria_selecionada': categoria_id
    })

@login_required
@acesso_restrito(['admin'])
def subcategoria_update(request, pk):
    sub = get_object_or_404(SubCategoriaDespesa, pk=pk)
    form = SubCategoriaDespesaForm(request.POST or None, instance=sub)

    if form.is_valid():
        form.save()
        messages.success(request, 'Subcategoria atualizada.')
        return redirect('subcategoria_create')

    return render(request, 'financeiro/subcategoria_form.html', {
        'form': form,
        'subcategorias': SubCategoriaDespesa.objects.select_related('categoria'),
        'categorias': CategoriaDespesa.objects.filter(ativa=True),
        'editando': True
    })

@login_required
@acesso_restrito(['admin'])
def subcategoria_delete(request, pk):
    sub = get_object_or_404(SubCategoriaDespesa, pk=pk)
    sub.delete()
    messages.success(request, 'Subcategoria removida.')
    return redirect('subcategoria_create')


# ============================
# AJAX
# ============================
@login_required
@acesso_restrito(['admin'])
def subcategorias_por_categoria(request):
    categoria_id = request.GET.get('categoria_id')
    subcategorias = SubCategoriaDespesa.objects.filter(
        categoria_id=categoria_id, ativa=True
    ).values('id', 'nome')

    return JsonResponse(list(subcategorias), safe=False)


from django.http import JsonResponse
from autocarros.models import SubCategoriaDespesa  # ✅ IMPORT CERTO

@login_required
@acesso_restrito(['admin'])
def ajax_subcategorias(request):
    categoria_id = request.GET.get("categoria_id")

    qs = SubCategoriaDespesa.objects.filter(
        categoria_id=categoria_id
    ).order_by("nome")

    data = [
        {"id": s.id, "label": s.nome}
        for s in qs
    ]

    return JsonResponse(data, safe=False)


# -------- DESPESAS --------
@login_required
@acesso_restrito(['admin'])
def despesa_list(request):
    def to_int(valor, padrao):
        try:
            return int(str(valor).replace(".", ""))
        except (TypeError, ValueError):
            return padrao

    mes = to_int(request.GET.get("mes"), now().month)
    ano = to_int(request.GET.get("ano"), now().year)
    categoria_id = request.GET.get("categoria")

    despesas = (
        Despesa2.objects
        .filter(data__month=mes, data__year=ano)
        .select_related("categoria", "subcategoria")
    )

    if categoria_id:
        despesas = despesas.filter(categoria_id=categoria_id)

    despesas = despesas.order_by("data")

    total = despesas.aggregate(total=Sum("valor"))["total"] or 0

    totais_categoria = (
        despesas
        .values("categoria__nome")
        .annotate(total=Sum("valor"))
        .order_by("categoria__nome")
    )

    totais_subcategoria = (
        despesas
        .values("subcategoria__nome")
        .annotate(total=Sum("valor"))
        .order_by("subcategoria__nome")
    )

    categorias = CategoriaDespesa.objects.all().order_by("nome")

    context = {
        "despesas": despesas,
        "categorias": categorias,
        "mes": mes,
        "ano": ano,
        "categoria_id": categoria_id,
        "total": total,
        "totais_categoria": totais_categoria,
        "totais_subcategoria": totais_subcategoria,
    }

    return render(request, "financeiro/despesa_list.html", context)


@login_required
@acesso_restrito(['admin'])
def despesa_create(request):
    if request.method == "POST":
        form = DespesaForm2(request.POST)
        if form.is_valid():
            form.save()
            return redirect("despesa_list")
    else:
        form = DespesaForm2()

    return render(request, "financeiro/despesa_create.html", {"form": form})


@login_required
@acesso_restrito(['admin'])
def despesa_editar(request, pk):
    despesa = get_object_or_404(Despesa2, pk=pk)

    if request.method == "POST":
        form = DespesaForm2(request.POST, instance=despesa)
        if form.is_valid():
            form.save()
            messages.success(request, "Despesa atualizada com sucesso.")
            return redirect("despesa_list")
    else:
        form = DespesaForm2(instance=despesa)

    return render(request, "financeiro/despesa_create.html", {
        "form": form,
        "titulo": "Editar Despesa"
    })


@login_required
@acesso_restrito(['admin'])
def despesa_eliminar(request, pk):
    despesa = get_object_or_404(Despesa2, pk=pk)
    despesa.delete()
    return redirect("despesa_list")

